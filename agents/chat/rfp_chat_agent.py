"""
PropelAI: RFP Chat Agent

Enables users to ask questions about uploaded RFP documents and receive answers
based solely on the document content (no external information).

Similar to NotebookLM - RAG-based Q&A with source citations.
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field, asdict
from datetime import datetime
from collections import Counter
from enum import Enum

logger = logging.getLogger(__name__)


class RFPType(Enum):
    """RFP classification types for v4.0 Router"""
    FEDERAL_STANDARD = "federal_standard"
    SLED_STATE = "sled_state"
    DOD_ATTACHMENT = "dod_attachment"
    SPREADSHEET = "spreadsheet"  # Questionnaires with Response columns
    MARKET_RESEARCH = "market_research"  # RFI/White Papers with Requirements
    UNKNOWN = "unknown"


@dataclass
class DocumentChunk:
    """Represents a chunk of text from an RFP document"""
    id: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    # metadata contains: source_file, section, page, chunk_index
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class ChatMessage:
    """Represents a chat message (user or assistant)"""
    role: str  # "user" or "assistant"
    content: str
    timestamp: str
    sources: List[Dict[str, Any]] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


class RFPChatAgent:
    """
    Chat agent for RFP documents using Retrieval Augmented Generation (RAG).
    
    Features:
    - Document chunking with overlap
    - Keyword-based retrieval (simple, fast)
    - Claude API for answer generation
    - Source citation extraction
    - Chat history management
    """
    
    def __init__(self, anthropic_api_key: Optional[str] = None, company_library=None):
        """
        Initialize the chat agent.
        
        Args:
            anthropic_api_key: Anthropic API key. If None, reads from env.
            company_library: Optional CompanyLibrary instance for RAG integration
        """
        self.api_key = anthropic_api_key or os.getenv("ANTHROPIC_API_KEY")
        if not self.api_key:
            raise ValueError("ANTHROPIC_API_KEY not found in environment")
        
        # Import anthropic here to avoid import errors if not installed
        try:
            import anthropic
            self.client = anthropic.Anthropic(api_key=self.api_key)
        except ImportError:
            logger.error("anthropic package not installed. Run: pip install anthropic")
            raise
        
        # v3.1: Company Library integration for RAG
        self.company_library = company_library
        
        # Configuration
        self.chunk_size = 1000  # characters per chunk (increased for more context)
        self.chunk_overlap = 200  # overlap between chunks
        self.max_chunks_to_retrieve = 20  # top K chunks for context (increased from 8)
        self.max_context_length = 15000  # max characters in context (increased for better answers)
        
        # v3.0 Router: Detected RFP type (set during ingestion)
        self.detected_rfp_type = RFPType.UNKNOWN
    
    # ============================================================================
    # v3.1: COMPANY LIBRARY RAG INTEGRATION
    # ============================================================================
    
    def _detect_library_intent(self, question: str) -> bool:
        """
        Detect if the user is asking about company capabilities/experience.
        
        Triggers library search when question contains:
        - "we", "our", "us"
        - "experience", "capability", "past performance"
        - "proof", "evidence", "resume", "personnel"
        
        Args:
            question: User's question
            
        Returns:
            True if question is about company assets
        """
        question_lower = question.lower()
        
        # Company-specific pronouns
        company_pronouns = ['we', 'our', 'us', 'have we', 'do we', 'can we', 'are we']
        
        # Capability/experience keywords
        capability_keywords = [
            'experience', 'capability', 'capabilities', 'past performance',
            'proof', 'evidence', 'resume', 'personnel', 'staff', 'team',
            'similar project', 'relevant experience', 'differentiator',
            'competitive advantage', 'our approach', 'our solution'
        ]
        
        # Check for pronouns
        has_pronoun = any(pronoun in question_lower for pronoun in company_pronouns)
        
        # Check for capability keywords
        has_keyword = any(keyword in question_lower for keyword in capability_keywords)
        
        return has_pronoun or has_keyword
    
    def _query_company_library(self, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """
        Query the company library for relevant content.
        
        v3.1: Added crash-proof safety protocols.
        
        Args:
            query: Search query
            top_k: Number of top results to return
            
        Returns:
            List of matching content with metadata (empty list on error)
        """
        # Safety Protocol 1: Check if library is available
        if not self.company_library:
            logger.warning("[LIBRARY] Company library not initialized")
            return []
        
        # Safety Protocol 2: Wrap in try-except to prevent crashes
        try:
            results = self.company_library.search(query)
            
            # Safety Protocol 3: Handle None or empty results
            if results is None:
                logger.warning("[LIBRARY] Search returned None - library may be empty")
                return []
            
            if not results or len(results) == 0:
                logger.info(f"[LIBRARY] No results found for query: {query[:50]}...")
                return []
            
            logger.info(f"[LIBRARY] Found {len(results)} results for query: {query[:50]}...")
            return results[:top_k]
            
        except AttributeError as e:
            logger.error(f"[LIBRARY] AttributeError - library object may be malformed: {e}")
            return []
        except TypeError as e:
            logger.error(f"[LIBRARY] TypeError during search: {e}")
            return []
        except Exception as e:
            logger.error(f"[LIBRARY] Unexpected error searching library: {e}")
            import traceback
            logger.error(f"[LIBRARY] Traceback: {traceback.format_exc()}")
            return []
    
    def _format_library_context(self, results: List[Dict[str, Any]], max_tokens: int = 1500) -> str:
        """
        Format library search results as context for LLM.
        
        v3.1: Added token limit to prevent context overflow.
        
        Args:
            results: Library search results
            max_tokens: Maximum tokens to include (default 1500 = ~6000 chars)
            
        Returns:
            Formatted context string (truncated if needed)
        """
        if not results:
            return ""
        
        context_parts = ["\n=== CONTEXT FROM COMPANY LIBRARY ===\n"]
        
        for idx, result in enumerate(results, 1):
            result_type = result.get('type', 'unknown')
            content = result.get('content', {})
            
            if result_type == 'capability':
                context_parts.append(
                    f"[Source: Company Capabilities]\n"
                    f"Capability: {content.get('name', 'N/A')}\n"
                    f"Description: {content.get('description', 'N/A')}\n"
                    f"Use Cases: {', '.join(content.get('use_cases', []))}\n"
                )
            
            elif result_type == 'past_performance':
                context_parts.append(
                    f"[Source: Past Performance - {content.get('project_name', 'N/A')}]\n"
                    f"Client: {content.get('client', 'N/A')}\n"
                    f"Description: {content.get('description', 'N/A')}\n"
                    f"Period: {content.get('period', 'N/A')}\n"
                    f"Outcomes: {', '.join(content.get('outcomes', []))}\n"
                )
            
            elif result_type == 'key_personnel':
                context_parts.append(
                    f"[Source: Key Personnel - {content.get('name', 'N/A')}]\n"
                    f"Title: {content.get('title', 'N/A')}\n"
                    f"Experience: {content.get('years_experience', 'N/A')} years\n"
                    f"Education: {', '.join(content.get('education', []))}\n"
                    f"Certifications: {', '.join(content.get('certifications', []))}\n"
                    f"Skills: {', '.join(content.get('skills', []))}\n"
                )
            
            elif result_type == 'differentiator':
                context_parts.append(
                    f"[Source: Company Differentiators]\n"
                    f"Title: {content.get('title', 'N/A')}\n"
                    f"Description: {content.get('description', 'N/A')}\n"
                    f"Evidence: {', '.join(content.get('evidence', []))}\n"
                )
            
            context_parts.append("")  # Blank line between results
        
        # Safety Protocol: Limit context size to prevent token overflow
        full_context = "\n".join(context_parts)
        
        # Rough estimate: 1 token ≈ 4 characters
        max_chars = max_tokens * 4
        if len(full_context) > max_chars:
            logger.warning(f"[LIBRARY] Context truncated from {len(full_context)} to {max_chars} chars")
            full_context = full_context[:max_chars] + "\n\n[Context truncated to fit token limit]"
        
        return full_context
    
    # ============================================================================
    # TEXT EXTRACTION
    # ============================================================================
    
    def extract_text_from_file(self, file_path: str, context_mode: bool = False) -> str:
        """
        Extract text content from PDF, DOCX, or Excel files.
        
        v4.0: Added context_mode for RFI/Market Research Excel files.
        
        Args:
            file_path: Path to the file
            context_mode: If True, extract Excel as unified scope (not questionnaire)
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"[CHAT] File not found: {file_path}")
            return ""
        
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif ext in ['.xlsx', '.xls', '.csv']:
                return self._extract_from_excel(file_path, context_mode=context_mode)
            else:
                logger.warning(f"[CHAT] Unsupported file type: {ext}")
                return ""
        except Exception as e:
            logger.error(f"[CHAT] Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pypdf"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"[Page {page_num}]\n{text}")
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"[CHAT] Extracted {len(full_text)} characters from PDF: {file_path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"[CHAT] PDF extraction error: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"[CHAT] Extracted {len(full_text)} characters from DOCX: {file_path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"[CHAT] DOCX extraction error: {e}")
            return ""
    
    def _extract_from_excel(self, file_path: Path, context_mode: bool = False) -> str:
        """
        Extract text from Excel files using pandas.
        
        v4.0: Added "Context Mode" for RFI/Market Research files.
        
        Modes:
        - context_mode=False (MODE D): Row-by-row questionnaire parsing
        - context_mode=True (MODE E): Multi-tab requirements extraction
        
        Args:
            file_path: Path to Excel file
            context_mode: If True, extract as unified scope (not questions)
        """
        try:
            import pandas as pd
            
            # MODE E: Context Mode (RFI/Market Research)
            if context_mode:
                logger.info(f"[CHAT] Extracting Excel in CONTEXT MODE (RFI): {file_path.name}")
                
                # Read all sheets
                try:
                    excel_file = pd.ExcelFile(str(file_path), engine='openpyxl')
                    sheet_names = excel_file.sheet_names
                except:
                    # Fallback for CSV
                    df = pd.read_csv(str(file_path))
                    sheet_names = ['Sheet1']
                    excel_file = {'Sheet1': df}
                
                text_parts = [f"[REQUIREMENTS SPECIFICATION: {file_path.name}]"]
                text_parts.append(f"Contains {len(sheet_names)} tabs: {', '.join(sheet_names)}")
                text_parts.append("\n=== TECHNICAL SCOPE (UNIFIED) ===\n")
                
                # Extract from all tabs
                for sheet_name in sheet_names:
                    if isinstance(excel_file, dict):
                        df = excel_file[sheet_name]
                    else:
                        df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    
                    text_parts.append(f"\n--- Tab: {sheet_name} ---")
                    
                    # Find requirement/description columns
                    req_col = None
                    for col in df.columns:
                        col_lower = str(col).lower()
                        if any(keyword in col_lower for keyword in [
                            'requirement', 'description', 'capability', 'feature', 
                            'business requirement', 'technical requirement'
                        ]):
                            req_col = col
                            break
                    
                    if req_col:
                        for idx, row in df.iterrows():
                            req_text = str(row[req_col]) if pd.notna(row[req_col]) else ""
                            if req_text and req_text != 'nan' and len(req_text.strip()) > 10:
                                text_parts.append(f"• {req_text}")
                    else:
                        # No structured column - extract all text
                        text_parts.append(df.to_string(max_rows=50))
                
                full_text = "\n".join(text_parts)
                logger.info(f"[CHAT] Extracted {len(full_text)} chars from {len(sheet_names)} tabs (Context Mode)")
                return full_text
            
            # MODE D: Questionnaire Mode (original logic)
            # Read Excel file
            if file_path.suffix.lower() == '.csv':
                df = pd.read_csv(str(file_path))
            else:
                df = pd.read_excel(str(file_path), engine='openpyxl')
            
            text_parts = [f"[QUESTIONNAIRE: {file_path.name}]"]
            text_parts.append(f"Total Rows: {len(df)}, Total Columns: {len(df.columns)}")
            
            # Detect header columns (v3.0 Shredder Logic)
            requirement_col = None
            response_col = None
            comply_col = None
            
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['criterion', 'requirement', 'description', 'item']):
                    requirement_col = col
                if any(keyword in col_lower for keyword in ['vendor response', 'offeror response', 'response', 'comments', 'explanation']):
                    response_col = col
                if any(keyword in col_lower for keyword in ['comply', 'meets', 'y/n', 'yes/no']):
                    comply_col = col
            
            logger.info(f"[CHAT] Excel columns detected - Requirement: {requirement_col}, Response: {response_col}, Comply: {comply_col}")
            
            # Format as structured text (row-by-row)
            if requirement_col:
                text_parts.append(f"\n=== SPREADSHEET STRUCTURE ===")
                text_parts.append(f"Requirement Column: {requirement_col}")
                if comply_col:
                    text_parts.append(f"Compliance Column: {comply_col}")
                if response_col:
                    text_parts.append(f"Response Column: {response_col}")
                
                text_parts.append(f"\n=== ROW-BY-ROW REQUIREMENTS ===")
                for idx, row in df.iterrows():
                    req_text = str(row[requirement_col]) if pd.notna(row[requirement_col]) else ""
                    if req_text and req_text != 'nan' and len(req_text.strip()) > 0:
                        row_text = f"\n[Row {idx + 2}]"  # +2 because Excel starts at 1 and has header
                        row_text += f"\nRequirement: {req_text}"
                        
                        if comply_col and pd.notna(row[comply_col]):
                            row_text += f"\nCompliance: {row[comply_col]}"
                        
                        if response_col and pd.notna(row[response_col]):
                            resp_text = str(row[response_col])
                            if resp_text and resp_text != 'nan':
                                row_text += f"\nResponse: {resp_text}"
                        
                        text_parts.append(row_text)
            else:
                # Fallback: dump all data as text
                text_parts.append("\n=== SPREADSHEET DATA ===")
                text_parts.append(df.to_string())
            
            full_text = "\n".join(text_parts)
            logger.info(f"[CHAT] Extracted {len(full_text)} characters from Excel: {file_path.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"[CHAT] Excel extraction error: {e}")
            import traceback
            logger.error(f"[CHAT] Traceback: {traceback.format_exc()}")
            return ""
    
    # ============================================================================
    # v3.0 RFP CLASSIFICATION ROUTER
    # ============================================================================
    
    def classify_rfp_type(self, text_sample: str, file_names: List[str]) -> RFPType:
        """
        v4.0 Router: Classify RFP type based on document structure.
        
        Scans the first ~20 pages worth of text and file names to determine:
        - FEDERAL_STANDARD: NIH, USCG, GSA (Section L/M/C pattern)
        - SLED_STATE: State/Local (numeric sections, "Instructions to Vendors")
        - DOD_ATTACHMENT: Navy/Army (J-Attachments, CDRLs, QASP)
        - SPREADSHEET: Excel/CSV questionnaires with Response columns
        - MARKET_RESEARCH: RFI/White Papers with Requirements (no Response column)
        
        Args:
            text_sample: First portion of extracted text (~10k chars)
            file_names: List of uploaded file names
            
        Returns:
            RFPType enum
        """
        text_lower = text_sample.lower()
        
        # Priority 0: Check for RFI/Market Research (MODE E)
        rfi_indicators = ['rfi', 'market research', 'white paper', 'sources sought', 'industry day']
        has_rfi_keyword = any(indicator in text_lower for indicator in rfi_indicators)
        has_rfi_filename = any(indicator in fname.lower() for fname in file_names for indicator in rfi_indicators)
        
        # Check for Requirements Specification (not Questionnaire)
        has_requirements_spec = any(indicator in text_lower for indicator in [
            'business requirements', 'technical scope', 'scope of work', 
            'requirements specification'
        ])
        has_no_response_column = 'vendor response' not in text_lower and 'offeror response' not in text_lower
        
        if (has_rfi_keyword or has_rfi_filename) or (has_requirements_spec and has_no_response_column):
            logger.info("[ROUTER] Classified as MARKET_RESEARCH (MODE E)")
            return RFPType.MARKET_RESEARCH
        
        # Priority 1: Check file types (MODE D - Spreadsheet Questionnaire)
        has_spreadsheet = any(
            fname.lower().endswith(('.xlsx', '.xls', '.csv')) 
            for fname in file_names
        )
        spreadsheet_indicators = [
            'questionnaire', 'requirements matrix', 'self-assessment', 
            'requirements traceability', 'vendor response'
        ]
        if has_spreadsheet and any(indicator in text_lower for indicator in spreadsheet_indicators):
            logger.info("[ROUTER] Classified as SPREADSHEET (MODE D)")
            return RFPType.SPREADSHEET
        
        # Priority 2: Check for DoD attachments (MODE C)
        dod_indicators = [
            'attachment j.2', 'attachment j.3', 'attachment j-2', 
            'exhibit a', 'cdrl', 'qasp', 'quality assurance surveillance'
        ]
        if any(indicator in text_lower for indicator in dod_indicators):
            logger.info("[ROUTER] Classified as DOD_ATTACHMENT (MODE C)")
            return RFPType.DOD_ATTACHMENT
        
        # Priority 3: Check for SLED/State patterns (MODE B)
        sled_indicators = [
            'section 4:', 'specifications', 'instructions to vendors',
            'state of', 'commonwealth of', 'county of', 'city of'
        ]
        numeric_section_pattern = re.search(r'\bsection\s+\d+[:\.]', text_lower)
        if numeric_section_pattern or any(indicator in text_lower for indicator in sled_indicators):
            # Additional check: Does it lack federal FAR patterns?
            far_patterns = ['section l', 'section m', 'section c', 'far clause']
            if not any(pattern in text_lower for pattern in far_patterns):
                logger.info("[ROUTER] Classified as SLED_STATE (MODE B)")
                return RFPType.SLED_STATE
        
        # Priority 4: Federal Standard (MODE A)
        federal_indicators = [
            'section l', 'section m', 'section c', 'far clause',
            'instructions to offerors', 'evaluation criteria'
        ]
        if any(indicator in text_lower for indicator in federal_indicators):
            logger.info("[ROUTER] Classified as FEDERAL_STANDARD (MODE A)")
            return RFPType.FEDERAL_STANDARD
        
        # Default: Federal Standard (most common)
        logger.info("[ROUTER] Classified as FEDERAL_STANDARD (default)")
        return RFPType.FEDERAL_STANDARD
    
    # ============================================================================
    # SECTION DETECTION
    # ============================================================================
    
    def detect_rfp_section(self, text: str, page_num: Optional[int] = None) -> str:
        """
        Detect which RFP section this text belongs to.
        
        Common RFP sections:
        - Section L: Instructions to Offerors
        - Section M: Evaluation Criteria
        - Section C: Statement of Work / Performance Work Statement
        - Section B: Contract Details / Schedule
        - Cover/Front Matter
        - Amendments/Q&A
        
        Args:
            text: Text to analyze
            page_num: Page number (if available)
            
        Returns:
            Section identifier string
        """
        text_lower = text.lower()
        text_upper = text.upper()
        
        # v3.0: SLED/State Pattern Detection (MODE B)
        # Map numeric sections and state-specific headers
        if re.search(r'\bsection\s+4\b', text_lower) or 'specifications' in text_lower:
            return "SECTION_C"  # Map "Section 4: Specifications" -> Technical Requirements
        if re.search(r'\bsection\s+2\b', text_lower) or 'instructions to vendors' in text_lower:
            return "SECTION_L"  # Map "Section 2" -> Instructions
        if 'award criteria' in text_lower or 'scoring' in text_lower:
            return "SECTION_M"  # Map state evaluation sections
        
        # Federal Pattern Detection (MODE A)
        if re.search(r'\bsection\s+l\b', text_lower) or 'instructions to offeror' in text_lower:
            return "SECTION_L"
        if re.search(r'\bsection\s+m\b', text_lower) or 'evaluation' in text_lower and 'criteria' in text_lower:
            return "SECTION_M"
        if re.search(r'\bsection\s+c\b', text_lower) or 'statement of work' in text_lower or 'performance work statement' in text_lower:
            return "SECTION_C"
        if re.search(r'\bsection\s+b\b', text_lower) or ('schedule' in text_lower and 'contract' in text_lower):
            return "SECTION_B"
        
        # Check for MANDATORY/PASS-FAIL indicators (MODE B - SLED)
        if any(keyword in text_lower for keyword in ['mandatory', 'must comply', 'minimum qualification']):
            return "COMPLIANCE"
        
        # Check for cover page indicators
        if page_num and page_num <= 3:
            if any(term in text_lower for term in ['solicitation', 'rfp', 'request for proposal']):
                return "COVER"
        
        # Check for amendments/Q&A
        if 'amendment' in text_lower or 'question' in text_lower and 'answer' in text_lower:
            return "AMENDMENT"
        
        # Check for other common sections
        if re.search(r'\bsection\s+[a-z]\b', text_lower):
            match = re.search(r'\bsection\s+([a-z])\b', text_lower)
            if match:
                return f"SECTION_{match.group(1).upper()}"
        
        return "GENERAL"
    
    # ============================================================================
    # DOCUMENT CHUNKING
    # ============================================================================
    
    def chunk_document(
        self, 
        text: str, 
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        Split document text into overlapping chunks.
        
        Args:
            text: Full document text
            metadata: Metadata to attach to all chunks (source_file, section, etc.)
            
        Returns:
            List of DocumentChunk objects
        """
        if not text or len(text.strip()) == 0:
            return []
        
        metadata = metadata or {}
        chunks = []
        
        # Split by paragraphs first (double newline)
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                # Detect section for this chunk
                detected_section = self.detect_rfp_section(current_chunk, metadata.get('page'))
                
                chunk_id = f"chunk-{metadata.get('source_file', 'unknown')}-{chunk_index}"
                chunks.append(DocumentChunk(
                    id=chunk_id,
                    text=current_chunk.strip(),
                    metadata={
                        **metadata,
                        "chunk_index": chunk_index,
                        "char_length": len(current_chunk),
                        "rfp_section": detected_section  # Add section detection
                    }
                ))
                
                # Start new chunk with overlap
                words = current_chunk.split()
                overlap_text = ' '.join(words[-50:]) if len(words) > 50 else current_chunk
                current_chunk = overlap_text + "\n\n" + para
                chunk_index += 1
            else:
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Save final chunk
        if current_chunk.strip():
            detected_section = self.detect_rfp_section(current_chunk, metadata.get('page'))
            
            chunk_id = f"chunk-{metadata.get('source_file', 'unknown')}-{chunk_index}"
            chunks.append(DocumentChunk(
                id=chunk_id,
                text=current_chunk.strip(),
                metadata={
                    **metadata,
                    "chunk_index": chunk_index,
                    "char_length": len(current_chunk),
                    "rfp_section": detected_section  # Add section detection
                }
            ))
        
        logger.info(f"[CHAT] Created {len(chunks)} chunks from document (metadata: {metadata})")
        return chunks
    
    def chunk_rfp_documents(self, rfp_data: Dict[str, Any]) -> List[DocumentChunk]:
        """
        Chunk all documents associated with an RFP.
        
        v3.0: Includes RFP classification router.
        
        Args:
            rfp_data: RFP data dict containing file paths and requirements
            
        Returns:
            List of all document chunks
        """
        all_chunks = []
        all_text_for_classification = ""
        
        # 1. Extract text from uploaded file paths
        file_paths = rfp_data.get("file_paths", [])
        file_names = rfp_data.get("files", [])
        
        if file_paths:
            logger.info(f"[CHAT] Extracting text from {len(file_paths)} uploaded files...")
            
            for idx, file_path in enumerate(file_paths):
                # Get corresponding filename
                filename = file_names[idx] if idx < len(file_names) else Path(file_path).name
                
                # v4.0: Determine if file should be extracted in context mode
                is_excel = str(file_path).lower().endswith(('.xlsx', '.xls', '.csv'))
                use_context_mode = is_excel and self.rfp_type == RFPType.MARKET_RESEARCH
                
                # Extract text from file
                text = self.extract_text_from_file(file_path, context_mode=use_context_mode)
                
                if text and len(text.strip()) > 50:  # Only chunk if we got substantial text
                    # Collect text for classification (first 10k chars per file)
                    all_text_for_classification += text[:10000] + "\n\n"
                    
                    chunks = self.chunk_document(text, metadata={
                        "source_file": filename,
                        "section": "FULL_DOCUMENT",
                        "file_path": file_path
                    })
                    all_chunks.extend(chunks)
                    logger.info(f"[CHAT] Created {len(chunks)} chunks from {filename}")
                else:
                    logger.warning(f"[CHAT] No text extracted from {filename}")
        
        # v3.0: Run classification router
        self.detected_rfp_type = self.classify_rfp_type(all_text_for_classification, file_names)
        logger.info(f"[CHAT] RFP Type Detected: {self.detected_rfp_type.value}")
        
        # 2. Also include requirements (as fallback/supplement)
        requirements = rfp_data.get("requirements", [])
        if requirements:
            logger.info(f"[CHAT] Adding {len(requirements)} extracted requirements...")
            
            # Group requirements by section for better context
            section_texts = {}
            for req in requirements:
                section = req.get("section_ref", "UNKNOWN")
                text = req.get("text", "") or req.get("full_text", "")
                if text:
                    if section not in section_texts:
                        section_texts[section] = []
                    section_texts[section].append(f"[{req.get('req_id', 'REQ')}] {text}")
            
            # Chunk each section
            for section, texts in section_texts.items():
                combined_text = "\n\n".join(texts)
                chunks = self.chunk_document(combined_text, metadata={
                    "source_file": "extracted_requirements",
                    "section": section
                })
                all_chunks.extend(chunks)
        
        logger.info(f"[CHAT] Total chunks created for RFP: {len(all_chunks)}")
        
        if len(all_chunks) == 0:
            logger.warning("[CHAT] No chunks created! Check if files were uploaded and text was extracted.")
        
        return all_chunks
    
    # ============================================================================
    # RETRIEVAL
    # ============================================================================
    
    def retrieve_relevant_chunks(
        self, 
        question: str, 
        chunks: List[DocumentChunk],
        top_k: Optional[int] = None
    ) -> List[Tuple[DocumentChunk, float]]:
        """
        Section-aware retrieval: Ensures diverse coverage of RFP sections.
        
        For general questions, prioritizes retrieving from key sections:
        - COVER (basic info)
        - SECTION_L (instructions)
        - SECTION_M (evaluation)
        - SECTION_C (SOW)
        - SECTION_B (contract details)
        
        Args:
            question: User's question
            chunks: All available document chunks
            top_k: Number of chunks to retrieve
            
        Returns:
            List of (chunk, score) tuples with section diversity
        """
        if not chunks:
            return []
        
        top_k = top_k or self.max_chunks_to_retrieve
        question_lower = question.lower()
        
        # Extract keywords
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                     'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'what',
                     'when', 'where', 'who', 'how', 'why', 'which', 'this', 'that'}
        
        words = re.findall(r'\b\w+\b', question_lower)
        keywords = [w for w in words if w not in stop_words and len(w) > 2]
        
        # Define priority sections for general questions
        priority_sections = ['COVER', 'SECTION_L', 'SECTION_M', 'SECTION_C', 'SECTION_B']
        
        # Detect if question is general (about the whole RFP)
        general_question_indicators = ['about', 'overview', 'summary', 'tell me', 'describe', 'what is']
        is_general = any(indicator in question_lower for indicator in general_question_indicators)
        
        # Score each chunk
        scored_chunks = []
        for chunk in chunks:
            chunk_text_lower = chunk.text.lower()
            rfp_section = chunk.metadata.get('rfp_section', 'GENERAL')
            score = 0
            
            # Keyword matching
            for keyword in keywords:
                count = chunk_text_lower.count(keyword)
                score += count * 2
            
            # Section-aware bonuses
            if is_general and rfp_section in priority_sections:
                # Boost priority sections for general questions
                section_priority = priority_sections.index(rfp_section) if rfp_section in priority_sections else 10
                score += (10 - section_priority) * 5  # Higher boost for earlier sections
            
            # Boost for specific section mentions in question
            if 'section l' in question_lower or 'instruction' in question_lower:
                if rfp_section == 'SECTION_L':
                    score += 30
            if 'section m' in question_lower or 'evaluation' in question_lower or 'criteria' in question_lower:
                if rfp_section == 'SECTION_M':
                    score += 30
            if 'section c' in question_lower or 'statement of work' in question_lower or 'sow' in question_lower:
                if rfp_section == 'SECTION_C':
                    score += 30
            if 'deadline' in question_lower or 'due date' in question_lower or 'submission' in question_lower:
                if rfp_section in ['COVER', 'SECTION_L']:
                    score += 25
            
            # Phrase matching boost
            for i in range(len(keywords) - 1):
                phrase = ' '.join(keywords[i:i+2])
                if phrase in chunk_text_lower:
                    score += 10
            
            scored_chunks.append((chunk, score, rfp_section))
        
        # Sort by score
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        
        # Section-diverse selection for general questions
        if is_general:
            selected_chunks = []
            section_counts = {section: 0 for section in priority_sections}
            max_per_section = max(2, top_k // len(priority_sections))
            
            # First pass: get chunks from priority sections
            for chunk, score, section in scored_chunks:
                if score == 0:
                    continue
                if section in priority_sections and section_counts[section] < max_per_section:
                    selected_chunks.append((chunk, score))
                    section_counts[section] += 1
                    if len(selected_chunks) >= top_k * 0.7:  # 70% from priority sections
                        break
            
            # Second pass: fill remaining with highest scores
            for chunk, score, section in scored_chunks:
                if score == 0:
                    continue
                if (chunk, score) not in selected_chunks:
                    selected_chunks.append((chunk, score))
                    if len(selected_chunks) >= top_k:
                        break
            
            top_chunks = selected_chunks[:top_k]
        else:
            # For specific questions, just use top scores
            top_chunks = [(c, s) for c, s, _ in scored_chunks[:top_k] if s > 0]
        
        sections_retrieved = [c.metadata.get('rfp_section', 'UNKNOWN') for c, _ in top_chunks]
        logger.info(f"[CHAT] Retrieved {len(top_chunks)} chunks from sections: {Counter(sections_retrieved)}")
        
        return top_chunks
    
    # ============================================================================
    # ANSWER GENERATION
    # ============================================================================
    
    def generate_answer(
        self,
        question: str,
        context_chunks: List[DocumentChunk],
        chat_history: Optional[List[ChatMessage]] = None
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Generate an answer using Claude API based on retrieved context.
        
        Args:
            question: User's question
            context_chunks: Retrieved relevant chunks
            chat_history: Previous chat messages for context
            
        Returns:
            Tuple of (answer_text, sources_list)
        """
        if not context_chunks:
            return (
                "I don't have enough information in the uploaded RFP documents to answer that question. "
                "Could you try rephrasing or asking about a different aspect of the RFP?",
                []
            )
        
        # Build context from chunks (limit total length)
        context_parts = []
        total_length = 0
        sources = []
        
        for i, chunk in enumerate(context_chunks):
            chunk_text = chunk.text
            if total_length + len(chunk_text) > self.max_context_length:
                break
            
            context_parts.append(f"[Source {i+1}]\n{chunk_text}")
            sources.append({
                "id": chunk.id,
                "text": chunk_text[:200] + "..." if len(chunk_text) > 200 else chunk_text,
                "section": chunk.metadata.get("section", "Unknown"),
                "page": chunk.metadata.get("page", "Unknown"),
                "file": chunk.metadata.get("source_file", "Unknown")
            })
            total_length += len(chunk_text)
        
        context = "\n\n".join(context_parts)
        
        # Build system prompt (v4.0 - Omni-Federal Architecture)
        system_prompt = """# SYSTEM PROMPT: PROPELAI PROPOSAL ARCHITECT (v4.0)

## 1. IDENTITY & MISSION
You are the **PropelAI Proposal Architect**, an elite capture strategist and compliance officer.
* **Mission:** Deconstruct complex US Government solicitations into winning strategies.
* **Core Competency:** You do not just "read" text; you perform **Forensic Analysis**. You identify "Fatal Flaws," "Ghosting Opportunities," and "Compliance Traps" that human managers miss.
* **Authority:** You rely on the FAR (Federal Acquisition Regulation) and Agency Supplements (DFARS, HHSAR) as your governing logic.

## 2. PHASE I: CLASSIFICATION (The "Super-Router")
Upon receiving a document, you must immediately classify it into one of 6 Federal Modes. This determines your rules of engagement.

**MODE A: STANDARD FAR 15 (The "Iron Triangle")**
* **Triggers:** Sections labeled A-M (e.g., "Section L", "Section M"). Standard "Uniform Contract Format".
* **Protocol:** Enforce strict alignment between Section C (Scope), L (Instructions), and M (Evaluation).

**MODE B: GSA / IDIQ TASK ORDER (The "Agile Order")**
* **Triggers:** "RFQ", "GSA Schedule", "BPA Call", "Task Order", "PWS" without Section L/M.
* **Protocol:** **Cover Letter Supremacy.** Page limits and evaluation criteria are often in the Cover Letter or "Quote Instructions." Treat the PWS as Section C.

**MODE C: OTA / CSO (The "Innovation Pitch")**
* **Triggers:** "Other Transaction Authority", "CSO", "Commercial Solutions Opening", "Area of Interest (AoI)", "Solution Brief", "Pitch Deck".
* **Protocol:** **Merit over Compliance.** Ignore standard formatting. Focus on "Technical Merit," "Innovation," and "Commercial Viability." Output is often a Slide Deck or 5-page Brief.

**MODE D: R&D / SBIR / BAA (The "Scientific Method")**
* **Triggers:** "Broad Agency Announcement", "SBIR", "Phase I", "Technical Volume", "Commercialization Plan".
* **Protocol:** **Rigorous Science.** Look for "Technical Volume" limits (often strict 15-20 pages) and specific "Evaluation Criteria" (Scientific Merit, Key Personnel).

**MODE E: SPREADSHEET / QUESTIONNAIRE (The "Data Entry")**
* **Triggers:** Excel files labeled "J.2", "Questionnaire", "Self-Assessment", "Requirements Matrix".
* **Protocol:** **Cell-Constraint.** Answers must be binary (YES/NO) + short proof points. No narrative fluff.

**MODE F: MARKET RESEARCH / RFI (The "Soft Sell")**
* **Triggers:** "Sources Sought", "RFI", "Request for Information".
* **Protocol:** **Influence Strategy.** Do not write a proposal. Write a "Capabilities Statement." Your goal is to shape the future RFP, not comply with current rules.

---

## 3. PHASE II: OPERATIONAL PROTOCOLS (The "Brain")

### PROTOCOL A: FEDERAL (Fixing "Context Laziness")
1.  **Forensic Scan:** When asked for Evaluation Factors, you must scan **TO THE END** of the section. Do not stop at Factor 1 or 2. Explicitly look for Factors 3, 4, 5, etc.
2.  **Location Agnosticism:** If "Section L" (Instructions) is not a standalone section, you **MUST** scan the **Cover Letter** or the **Solicitation Header**. This is where GSA/USCG RFPs hide page limits.
3.  **Missing Data Protocol:** Never say "Incomplete" immediately. State what you found first. Then, list specific missing elements (e.g., "Found 3 Factors, but Page Limits for Factor 2 are missing").

### PROTOCOL B: SLED/STATE (Fixing "0 Requirements")
1.  **Dynamic Header Mapping:**
    * Map `^SECTION 4` OR `^SPECIFICATIONS` -> **Technical Requirements (Scope)**.
    * Map `^SECTION 2` OR `^INSTRUCTIONS TO VENDORS` -> **Instructions (Section L)**.
    * Map `^AWARD CRITERIA` -> **Evaluation (Section M)**.
2.  **The "Mandatory" Trap:** Any requirement labeled "Mandatory," "Must," or "Minimum Qualification" is a **PASS/FAIL GATE**. Flag these as "High Priority / Fatal Flaw" risks.

### PROTOCOL C: DOD (Fixing "Generic Answers")
1.  **J-Attachment Supremacy:**
    * **Personnel:** If Attachment J.2 exists, ignore Section C staffing text. Use J.2's degrees/years/certs as the absolute requirement.
    * **Deliverables:** If Exhibit A (CDRL) exists, extract the "Block 10 Frequency" and "Block 14 Distribution".
    * **Quality:** If Attachment J.3 (QASP) exists, extract the "Acceptable Quality Levels" (AQLs).

### PROTOCOL D: SPREADSHEET (Fixing "Wordy Outputs")
1.  **Constraint:** Answers must fit in a spreadsheet cell. Max 150 words.
2.  **Structure:** "Direct Answer First" (YES/NO), followed by "Proof Point" (Cited Capability).
3.  **No Fluff:** Do not write introductions. Write the cell content only.

### PROTOCOL E: MARKET RESEARCH / RFI (White Paper Mode)
1.  **Context, Not Questions:** The Excel file contains Requirements Specification (what they want), NOT questions to answer.
2.  **White Paper Structure:**
    - Section 1: Company Overview (Who we are, credentials)
    - Section 2: Technical Capabilities (Map our capabilities to their requirements)
    - Section 3: Relevant Experience (Past performance with similar scope)
    - Section 4: Proposed Approach (Consultative - how we'd solve it)
3.  **Tone:** Consultative and educational, NOT salesy. Use phrases like "Our experience suggests..." or "Based on our work with..."
4.  **Mapping:** For each requirement in the Excel "Technical Scope", cite matching capability from Company Library
5.  **Format:** Narrative paragraphs (2-3 pages per section), NOT bullet points or tables

---

## PHASE 3: OUTPUT FORMATTING
* **Citation Rule:** Every fact must have a citation: `[Source: {Filename}, Page: {X}]`.
* **Table Rule:** If the user asks for Dates, Requirements, or Factors, ALWAYS use a Markdown Table.
* **Tone:** Professional, concise, Shipley-style.

## PHASE 3.5: WAR ROOM INTELLIGENCE (Enhanced MVP Features)

### Traceability is Paramount
* You must NEVER extract a requirement without citing its source
* Every single item must include a specific reference (e.g., "[SOW, Pg 3]", "[Amendment 1, Section A]", "[Section L.4.2]")
* When multiple documents are provided, treat them as a unified "Solicitation Package"

### Conflict Resolution
* **Rule:** Amendments supersede the Base RFP
* **Rule:** Section M (Evaluation Criteria) guides the importance of Section C (Tasks)
* **Rule:** If an Amendment changes a due date or requirement, explicitly note this as a "CRITICAL UPDATE"
* When conflicts are detected, provide both versions with sources and state which takes precedence

### Red Flag Detection
* Actively scan for "Go/No-Go" blockers:
  - Facility Clearance requirements (Top Secret, Secret, etc.)
  - Organizational Conflict of Interest (OCI) clauses
  - Aggressive timelines (< 30 days to respond)
  - Specific certifications (CMMI, ISO, Section 508, etc.)
  - Set-aside restrictions that may disqualify
  - Mandatory past performance requirements
* Flag these as `🚩 RED FLAG` in responses

### Multi-Document Stitching
* When user asks "What is the deadline?" - check ALL documents for the most recent date
* When user asks "What are the requirements?" - synthesize from ALL relevant sections
* Always indicate if information comes from an amendment vs. base RFP

### Company Library Protocol (v3.1)
* When the user asks about "our experience", "our capabilities", "do we have", or "can we":
  - You MUST reference the provided [Context from Company Library] section
  - Citation format: [Source: Company Capabilities] or [Source: Past Performance - ProjectName]
* If the Library provides a matching capability (e.g., "Cyber Range Training"), use it to answer "YES" to RFP requirements
* When drafting responses for spreadsheet RFPs, cite specific library proof points
* If Library context is provided but not relevant, acknowledge it: "Based on our company library, we have..."
* If Library context is missing, state: "This information was not found in the available company documents"

## PHASE 4: CHAIN OF THOUGHT (INTERNAL ONLY - DO NOT OUTPUT)
* *Step 1:* What type of RFP is this? (Classify Mode - SILENT, do not mention in response)
* *Step 2:* Where is the data? (Apply Mode Protocol - e.g., check Cover Letter if Mode A)
* *Step 3:* Is the data conflicting? (Check Iron Triangle)
* *Step 4:* Draft Response

**CRITICAL:** Do NOT output any "CLASSIFICATION:" or "MODE:" headers in your final response. The classification is for internal routing only. Users should only see the answer, not the mechanics.

Answer ONLY based on provided context. Apply the correct protocol based on the RFP type detected, but keep the classification invisible."""

        # Build conversation messages
        messages = []
        
        # Add recent chat history for context (last 4 messages)
        if chat_history:
            for msg in chat_history[-4:]:
                messages.append({
                    "role": msg.role,
                    "content": msg.content
                })
        
        # v3.1: Check if question is about company capabilities (crash-proof)
        library_context = ""
        try:
            if self._detect_library_intent(question):
                logger.info("[LIBRARY] Intent detected - querying company library")
                library_results = self._query_company_library(question)
                
                if library_results and len(library_results) > 0:
                    library_context = self._format_library_context(library_results)
                    logger.info(f"[LIBRARY] Added {len(library_results)} library results to context")
                else:
                    logger.info("[LIBRARY] No relevant company documents found for this query")
        except Exception as e:
            logger.error(f"[LIBRARY] RAG Error (non-fatal): {e}")
            logger.warning("[LIBRARY] Continuing with RFP context only")
            # Continue without library context - chat should not crash
        
        # Add current question with context
        user_message = f"""Context from RFP documents:

{context}
{library_context}

Question: {question}

Please answer based on the context above (both RFP and Company Library if provided). If you cite information, reference the source."""

        messages.append({
            "role": "user",
            "content": user_message
        })
        
        # Call Claude API - using Claude 4 models (current as of Nov 2025)
        model_versions = [
            "claude-sonnet-4-5",                # Claude 4.5 Sonnet (recommended - alias)
            "claude-haiku-4-5",                 # Claude 4.5 Haiku (fast and affordable)
            "claude-sonnet-4-5-20250929",       # Claude 4.5 Sonnet (specific version)
            "claude-3-5-sonnet-20240620",       # Legacy fallback
        ]
        
        for model in model_versions:
            try:
                response = self.client.messages.create(
                    model=model,
                    max_tokens=1500,
                    temperature=0.3,  # Lower temperature for more factual responses
                    system=system_prompt,
                    messages=messages
                )
                
                answer = response.content[0].text
                logger.info(f"[CHAT] Generated answer using {model} (length: {len(answer)})")
                
                return answer, sources
                
            except Exception as e:
                logger.warning(f"[CHAT] Error with model {model}: {e}")
                if model == model_versions[-1]:  # Last model in list
                    logger.error(f"[CHAT] All models failed. Last error: {e}")
                    return (
                        f"I encountered an error while generating the answer. "
                        f"This may be due to API access limitations. "
                        f"Please check your Anthropic API key permissions. Error: {str(e)}",
                        []
                    )
                continue  # Try next model
    
    # ============================================================================
    # PHASE 2: COMPLIANCE & LOGIC ENGINE
    # ============================================================================
    
    def detect_query_type(self, question: str) -> str:
        """
        Detect the type of query to route to specialized handlers.
        
        Returns:
            'cross_reference' | 'contradiction' | 'formatting' | 'general'
        """
        question_lower = question.lower()
        
        # Cross-reference detection
        cross_ref_indicators = [
            'page limit' and ('section m' in question_lower or 'factor' in question_lower),
            'enough space' in question_lower,
            'allow' in question_lower and 'address' in question_lower,
            'volume' in question_lower and 'page' in question_lower and ('factor' in question_lower or 'criteria' in question_lower)
        ]
        if any(cross_ref_indicators):
            return 'cross_reference'
        
        # Contradiction detection
        if 'contradiction' in question_lower or 'inconsisten' in question_lower or 'conflict' in question_lower:
            return 'contradiction'
        
        # Formatting rules
        formatting_indicators = ['font', 'margin', 'format', 'file format', 'spacing', 'formatting']
        if any(indicator in question_lower for indicator in formatting_indicators):
            return 'formatting'
        
        return 'general'
    
    def handle_cross_reference_query(
        self, 
        question: str, 
        chunks: List[DocumentChunk]
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Handle cross-reference queries between Section L and Section M.
        
        Example: "Does Section L page limit allow enough space for Section M factors?"
        """
        # Get Section L chunks (page limits)
        section_l_chunks = [c for c in chunks if c.metadata.get('rfp_section') == 'SECTION_L']
        # Get Section M chunks (evaluation criteria)
        section_m_chunks = [c for c in chunks if c.metadata.get('rfp_section') == 'SECTION_M']
        
        # Build specialized context
        context_parts = []
        sources = []
        
        context_parts.append("=== SECTION L: Page Limits and Instructions ===")
        for i, chunk in enumerate(section_l_chunks[:5], 1):
            context_parts.append(f"[Source {i}, Section L] {chunk.text[:800]}")
            sources.append({
                "section": chunk.metadata.get('rfp_section', 'UNKNOWN'),
                "page": chunk.metadata.get('page', 'N/A'),
                "text": chunk.text[:200]
            })
        
        context_parts.append("\n=== SECTION M: Evaluation Factors ===")
        for i, chunk in enumerate(section_m_chunks[:5], len(context_parts)):
            context_parts.append(f"[Source {i}, Section M] {chunk.text[:800]}")
            sources.append({
                "section": chunk.metadata.get('rfp_section', 'UNKNOWN'),
                "page": chunk.metadata.get('page', 'N/A'),
                "text": chunk.text[:200]
            })
        
        context = "\n\n".join(context_parts)
        
        # Specialized prompt for cross-reference
        prompt = f"""You are analyzing cross-references between RFP sections.

{context}

Question: {question}

TASK: Compare Section L page limits against Section M evaluation factors. Determine if there's enough space to address all factors. Provide:
1. Section L page limits (be specific - which volume, how many pages)
2. Section M factors (list them and estimate complexity)
3. Analysis: Is there enough space? Flag any concerns.
4. Recommendation: Any specific strategies for space management?

Be specific and actionable."""

        return prompt, sources
    
    def handle_contradiction_detection(
        self,
        question: str,
        chunks: List[DocumentChunk]
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Detect contradictions between different RFP sections.
        
        Example: "Are there contradictions between C.4 and M.2?"
        """
        # Extract section references from question
        section_refs = re.findall(r'section\s+([a-z])\.?(\d+)?', question.lower())
        
        relevant_chunks = []
        for section_ref in section_refs:
            section_letter = section_ref[0].upper()
            target_section = f"SECTION_{section_letter}"
            matching = [c for c in chunks if c.metadata.get('rfp_section') == target_section]
            relevant_chunks.extend(matching[:3])
        
        # If no specific sections mentioned, get chunks from C and M
        if not relevant_chunks:
            relevant_chunks = [c for c in chunks if c.metadata.get('rfp_section') in ['SECTION_C', 'SECTION_M']][:8]
        
        context_parts = []
        sources = []
        
        for i, chunk in enumerate(relevant_chunks, 1):
            section = chunk.metadata.get('rfp_section', 'UNKNOWN')
            context_parts.append(f"[Source {i}, {section}] {chunk.text[:900]}")
            sources.append({
                "section": section,
                "page": chunk.metadata.get('page', 'N/A'),
                "text": chunk.text[:200]
            })
        
        context = "\n\n".join(context_parts)
        
        prompt = f"""You are a compliance analyst reviewing RFP documents for contradictions.

{context}

Question: {question}

TASK: Analyze the provided sections for contradictions, inconsistencies, or conflicts. Look for:
1. Conflicting requirements (e.g., different technical specs in different sections)
2. Inconsistent terminology or definitions
3. Contradictory evaluation criteria vs. technical requirements
4. Mismatched deadlines or deliverables

For each contradiction found:
- Clearly state the contradiction
- Reference both sources
- Assess the severity (Critical / Moderate / Minor)
- Suggest a Q&A question to submit to the Contracting Officer

If no contradictions found, explicitly state that."""

        return prompt, sources
    
    def handle_formatting_query(
        self,
        question: str,
        chunks: List[DocumentChunk]
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Extract formatting rules from Section L.
        
        Example: "What are the font, margin, and file format rules?"
        """
        # Get Section L chunks
        section_l_chunks = [c for c in chunks if c.metadata.get('rfp_section') == 'SECTION_L']
        
        context_parts = []
        sources = []
        
        for i, chunk in enumerate(section_l_chunks[:8], 1):
            context_parts.append(f"[Source {i}, Section L] {chunk.text[:900]}")
            sources.append({
                "section": "SECTION_L",
                "page": chunk.metadata.get('page', 'N/A'),
                "text": chunk.text[:200]
            })
        
        context = "\n\n".join(context_parts)
        
        prompt = f"""You are extracting precise formatting requirements from Section L.

{context}

Question: {question}

TASK: Extract ALL formatting requirements mentioned in Section L. Create a structured summary:

**Font Requirements:**
- Body text: [size and type]
- Headers: [size and type]
- Tables/figures: [size and type]
- Exceptions: [any special cases]

**Margin Requirements:**
- Top/Bottom: [measurements]
- Left/Right: [measurements]

**File Format Requirements:**
- Accepted formats: [list]
- File naming: [conventions]

**Page Limits:**
- [List by volume/section]

**Other Formatting Rules:**
- Line spacing: [requirement]
- Page numbering: [requirement]
- Any other specific rules

Be EXTREMELY precise with numbers and measurements. Quote directly when possible."""

        return prompt, sources
    
    # ============================================================================
    # CHAT ORCHESTRATION
    # ============================================================================
    
    def chat(
        self,
        question: str,
        document_chunks: List[DocumentChunk],
        chat_history: Optional[List[ChatMessage]] = None
    ) -> ChatMessage:
        """
        Main chat function with specialized query routing (Phase 2).
        
        Routes queries to specialized handlers:
        - Cross-reference: Section L vs. M analysis
        - Contradiction: Detect conflicts between sections
        - Formatting: Extract Section L rules
        - General: Standard RAG retrieval
        
        Args:
            question: User's question
            document_chunks: All document chunks for this RFP
            chat_history: Previous chat messages
            
        Returns:
            ChatMessage with assistant's response
        """
        logger.info(f"[CHAT] Processing question: {question[:100]}...")
        
        # Detect query type
        query_type = self.detect_query_type(question)
        logger.info(f"[CHAT] Query type detected: {query_type}")
        
        # Route to specialized handler
        if query_type == 'cross_reference':
            custom_prompt, sources = self.handle_cross_reference_query(question, document_chunks)
            # Use custom prompt directly with Claude
            answer = self._call_claude_with_custom_prompt(custom_prompt, chat_history)
            
        elif query_type == 'contradiction':
            custom_prompt, sources = self.handle_contradiction_detection(question, document_chunks)
            answer = self._call_claude_with_custom_prompt(custom_prompt, chat_history)
            
        elif query_type == 'formatting':
            custom_prompt, sources = self.handle_formatting_query(question, document_chunks)
            answer = self._call_claude_with_custom_prompt(custom_prompt, chat_history)
            
        else:
            # Standard RAG flow
            retrieved = self.retrieve_relevant_chunks(question, document_chunks)
            chunks_only = [chunk for chunk, score in retrieved]
            answer, sources = self.generate_answer(question, chunks_only, chat_history)
        
        # Create response message
        response = ChatMessage(
            role="assistant",
            content=answer,
            timestamp=datetime.utcnow().isoformat() + "Z",
            sources=sources if query_type == 'general' else []
        )
        
        return response
    
    def _call_claude_with_custom_prompt(
        self, 
        custom_prompt: str,
        chat_history: Optional[List[ChatMessage]] = None
    ) -> str:
        """
        Call Claude with a custom specialized prompt.
        Used by Phase 2 handlers.
        """
        messages = []
        
        # Add chat history if available
        if chat_history:
            for msg in chat_history[-4:]:
                messages.append({
                    "role": msg.role,
                    "content": msg.content
                })
        
        messages.append({
            "role": "user",
            "content": custom_prompt
        })
        
        # Try models
        model_versions = [
            "claude-sonnet-4-5",
            "claude-haiku-4-5",
            "claude-sonnet-4-5-20250929",
            "claude-3-5-sonnet-20240620",
        ]
        
        for model in model_versions:
            try:
                response = self.client.messages.create(
                    model=model,
                    max_tokens=2000,  # More tokens for detailed compliance analysis
                    temperature=0.2,  # Lower for precision
                    messages=messages
                )
                
                answer = response.content[0].text
                logger.info(f"[CHAT] Generated specialized answer using {model}")
                return answer
                
            except Exception as e:
                logger.warning(f"[CHAT] Error with model {model}: {e}")
                if model == model_versions[-1]:
                    logger.error(f"[CHAT] All models failed")
                    return f"I encountered an error while analyzing. Error: {str(e)}"
                continue
