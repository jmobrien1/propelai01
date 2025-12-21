"""
Company Library - Document Parser and Indexer
Extracts structured data from company documents for proposal enhancement

Supported document types:
- Capabilities Statements
- Past Performance Narratives
- Resumes/CVs
- Technical Approaches
- Corporate Information

v1.0 - Initial implementation
"""

import os
import re
import json
import uuid
import tempfile
import subprocess
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field, asdict
from enum import Enum
from pathlib import Path
from datetime import datetime


class DocumentType(str, Enum):
    """Types of company documents"""
    CAPABILITIES = "capabilities"
    PAST_PERFORMANCE = "past_performance"
    RESUME = "resume"
    TECHNICAL_APPROACH = "technical_approach"
    CORPORATE_INFO = "corporate_info"
    OTHER = "other"


@dataclass
class Capability:
    """A specific capability or service offering"""
    id: str
    name: str
    description: str
    category: str = ""
    keywords: List[str] = field(default_factory=list)
    use_cases: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class PastPerformance:
    """Past performance record"""
    id: str
    project_name: str
    client: str
    description: str
    contract_value: Optional[str] = None
    period: Optional[str] = None
    relevance: List[str] = field(default_factory=list)
    outcomes: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class KeyPersonnel:
    """Key personnel/resume information"""
    id: str
    name: str
    title: str
    summary: str
    years_experience: Optional[int] = None
    education: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    skills: List[str] = field(default_factory=list)
    relevant_experience: List[Dict] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class Differentiator:
    """Company differentiator or competitive advantage"""
    id: str
    title: str
    description: str
    evidence: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return asdict(self)


@dataclass
class CompanyProfile:
    """Overall company profile extracted from documents"""
    company_name: str = ""
    tagline: str = ""
    executive_summary: str = ""
    capabilities: List[Capability] = field(default_factory=list)
    past_performance: List[PastPerformance] = field(default_factory=list)
    key_personnel: List[KeyPersonnel] = field(default_factory=list)
    differentiators: List[Differentiator] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    contract_vehicles: List[str] = field(default_factory=list)
    naics_codes: List[str] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    
    def to_dict(self) -> Dict:
        return {
            "company_name": self.company_name,
            "tagline": self.tagline,
            "executive_summary": self.executive_summary,
            "capabilities": [c.to_dict() for c in self.capabilities],
            "past_performance": [p.to_dict() for p in self.past_performance],
            "key_personnel": [k.to_dict() for k in self.key_personnel],
            "differentiators": [d.to_dict() for d in self.differentiators],
            "certifications": self.certifications,
            "contract_vehicles": self.contract_vehicles,
            "naics_codes": self.naics_codes,
            "keywords": self.keywords,
        }


@dataclass
class ParsedDocument:
    """Result of parsing a company document"""
    id: str
    filename: str
    document_type: DocumentType
    title: str
    raw_text: str
    sections: List[Dict]
    extracted_data: Dict
    parse_date: str
    
    def to_dict(self) -> Dict:
        return {
            "id": self.id,
            "filename": self.filename,
            "document_type": self.document_type.value,
            "title": self.title,
            "sections": self.sections,
            "extracted_data": self.extracted_data,
            "parse_date": self.parse_date,
        }


class CompanyLibraryParser:
    """
    Parser for company documents - extracts structured data for proposal enhancement
    """
    
    def __init__(self):
        # Patterns for detecting document type
        self.type_patterns = {
            DocumentType.CAPABILITIES: [
                r"capabilities?\s+(?:statement|for|overview)",
                r"corporate\s+capabilities",
                r"service\s+offerings?",
                r"core\s+competenc",
                r"what\s+we\s+do",
                r"capabilities\s+for\s+government",
                r"#\s+.+capabilities",  # Title contains "capabilities"
            ],
            DocumentType.PAST_PERFORMANCE: [
                r"past\s+performance",
                r"project\s+experience",
                r"case\s+stud",
                r"client\s+success",
                r"relevant\s+experience",
            ],
            DocumentType.RESUME: [
                r"curriculum\s+vitae",
                r"resume",
                r"professional\s+experience",
                r"education\s+and\s+training",
                r"work\s+history",
            ],
            DocumentType.TECHNICAL_APPROACH: [
                r"technical\s+approach",
                r"solution\s+overview",
                r"methodology",
                r"our\s+approach",
            ],
            DocumentType.CORPORATE_INFO: [
                r"duns\s+number",
                r"cage\s+code",
                r"naics\s+codes?",
                r"sam\.gov",
                r"corporate\s+information",
            ],
        }
        
        # Section header patterns
        self.section_patterns = [
            r"^#+\s+(.+)$",  # Markdown headers
            r"^(\d+\.?\d*)\s+(.+)$",  # Numbered sections
            r"^([A-Z][A-Z\s]+):?\s*$",  # ALL CAPS headers
            r"^\*\*(.+)\*\*\s*$",  # Bold text as headers
        ]
        
        # Capability extraction patterns
        self.capability_patterns = [
            r"\*\*([^*]+)\*\*[:\s]+([^*\n]+)",  # **Title**: Description
            r"[-•]\s+\*\*([^*]+)\*\*[:\s]*(.+)",  # - **Title**: Description
        ]
        
    def parse_document(self, file_path: str, document_type: Optional[DocumentType] = None) -> ParsedDocument:
        """
        Parse a document and extract structured data
        
        Args:
            file_path: Path to document file
            document_type: Optional type hint
            
        Returns:
            ParsedDocument with extracted data
        """
        # Convert to markdown for text extraction
        raw_text = self._extract_text(file_path)
        
        # Detect document type if not specified
        if document_type is None:
            document_type = self._detect_document_type(raw_text)
        
        # Extract sections
        sections = self._extract_sections(raw_text)
        
        # Extract type-specific data
        extracted_data = self._extract_data(raw_text, sections, document_type)
        
        # Get title
        title = self._extract_title(raw_text, sections)
        
        return ParsedDocument(
            id=str(uuid.uuid4())[:8].upper(),
            filename=os.path.basename(file_path),
            document_type=document_type,
            title=title,
            raw_text=raw_text,
            sections=sections,
            extracted_data=extracted_data,
            parse_date=datetime.now().isoformat(),
        )
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from document using pandoc or python-docx fallback"""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".md":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".docx":
            # Try python-docx first (more reliable), then pandoc
            try:
                from docx import Document
                doc = Document(file_path)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)
                return "\n\n".join(paragraphs)
            except ImportError:
                # Fall back to pandoc
                return self._extract_with_pandoc(file_path)
        elif ext in [".doc", ".pdf", ".rtf"]:
            return self._extract_with_pandoc(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _extract_with_pandoc(self, file_path: str) -> str:
        """Extract text using pandoc"""
        try:
            result = subprocess.run(
                ["pandoc", file_path, "-t", "markdown", "--wrap=none"],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0:
                return result.stdout
            else:
                raise ValueError(f"Pandoc error: {result.stderr}")
        except FileNotFoundError:
            raise ValueError("Pandoc not installed - required for .doc/.pdf/.rtf parsing")
    
    def _detect_document_type(self, text: str) -> DocumentType:
        """Detect document type from content"""
        text_lower = text.lower()
        
        scores = {}
        for doc_type, patterns in self.type_patterns.items():
            score = sum(1 for p in patterns if re.search(p, text_lower))
            scores[doc_type] = score
        
        # Return highest scoring type, or OTHER if no matches
        best_type = max(scores, key=scores.get)
        if scores[best_type] > 0:
            return best_type
        return DocumentType.OTHER
    
    def _extract_sections(self, text: str) -> List[Dict]:
        """Extract document sections"""
        sections = []
        lines = text.split("\n")
        
        current_section = {"title": "Introduction", "level": 0, "content": [], "start_line": 0}
        
        for i, line in enumerate(lines):
            # Check for markdown headers
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                # Save previous section
                if current_section["content"]:
                    current_section["content"] = "\n".join(current_section["content"]).strip()
                    sections.append(current_section)
                
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = {
                    "title": title,
                    "level": level,
                    "content": [],
                    "start_line": i
                }
                continue
            
            # Check for numbered sections (e.g., "1.0 Introduction")
            numbered_match = re.match(r"^(\d+\.?\d*)\s+([A-Z][^.]+)$", line)
            if numbered_match:
                if current_section["content"]:
                    current_section["content"] = "\n".join(current_section["content"]).strip()
                    sections.append(current_section)
                
                current_section = {
                    "title": numbered_match.group(2).strip(),
                    "number": numbered_match.group(1),
                    "level": 1,
                    "content": [],
                    "start_line": i
                }
                continue
            
            # Add to current section content
            current_section["content"].append(line)
        
        # Don't forget the last section
        if current_section["content"]:
            current_section["content"] = "\n".join(current_section["content"]).strip()
            sections.append(current_section)
        
        return sections
    
    def _extract_title(self, text: str, sections: List[Dict]) -> str:
        """Extract document title"""
        lines = text.split("\n")
        
        # Look for H1 header
        for line in lines[:10]:
            match = re.match(r"^#\s+(.+)$", line)
            if match:
                return match.group(1).strip()
        
        # Use first section title
        if sections:
            return sections[0].get("title", "Untitled Document")
        
        # Use first non-empty line
        for line in lines[:5]:
            if line.strip():
                return line.strip()[:100]
        
        return "Untitled Document"
    
    def _extract_data(self, text: str, sections: List[Dict], doc_type: DocumentType) -> Dict:
        """Extract type-specific data from document"""
        
        if doc_type == DocumentType.CAPABILITIES:
            return self._extract_capabilities_data(text, sections)
        elif doc_type == DocumentType.PAST_PERFORMANCE:
            return self._extract_past_performance_data(text, sections)
        elif doc_type == DocumentType.RESUME:
            return self._extract_resume_data(text, sections)
        elif doc_type == DocumentType.TECHNICAL_APPROACH:
            return self._extract_technical_approach_data(text, sections)
        elif doc_type == DocumentType.CORPORATE_INFO:
            return self._extract_corporate_info_data(text, sections)
        else:
            return self._extract_generic_data(text, sections)
    
    def _extract_capabilities_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract capabilities statement data"""
        data = {
            "company_name": "",
            "executive_summary": "",
            "capabilities": [],
            "differentiators": [],
            "use_cases": [],
            "partnerships": [],
            "keywords": [],
            "offerings": [],  # Core service offerings
        }
        
        # Extract company name from title
        lines = text.split("\n")
        for line in lines[:5]:
            # Match "# Company Capabilities for X" or similar
            match = re.match(r"^#\s+(.+?)\s+[Cc]apabilities", line)
            if match:
                data["company_name"] = match.group(1).strip()
                break
            # Also match just "# Company Name" as H1
            elif line.startswith("# ") and "capabilities" not in line.lower():
                data["company_name"] = line[2:].strip()
                break
        
        # Find executive summary section
        for section in sections:
            title_lower = section["title"].lower()
            content = section.get("content", "")
            
            if "executive summary" in title_lower or ("overview" in title_lower and section.get("level", 0) <= 2):
                # Extract first substantive paragraph
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip() and len(p.strip()) > 100]
                if paragraphs:
                    data["executive_summary"] = paragraphs[0][:2000]
                break
        
        # If no executive summary found, use first substantial paragraph
        if not data["executive_summary"]:
            for section in sections[:3]:
                content = section.get("content", "")
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip() and len(p.strip()) > 100]
                if paragraphs:
                    data["executive_summary"] = paragraphs[0][:2000]
                    break
        
        # Extract capabilities/offerings from structured sections
        cap_id = 0
        for section in sections:
            title_lower = section["title"].lower()
            content = section.get("content", "")
            
            # Look for capability/offering sections
            if any(kw in title_lower for kw in ["capabilit", "offering", "solution", "service"]):
                # Extract bullet points with bold titles
                # Pattern: - **Title:** Description or -   **Title**: Description
                bold_bullets = re.findall(
                    r"[-•*]\s+(?:>\s*)?\*\*([^*]+)\*\*[:\s]*([^\n]*(?:\n(?![-•*])[^\n]+)*)",
                    content
                )
                for name, desc in bold_bullets:
                    cap_id += 1
                    cap = {
                        "id": f"CAP-{cap_id:03d}",
                        "name": name.strip()[:100],
                        "description": desc.strip()[:500] if desc.strip() else name.strip(),
                        "category": section["title"],
                    }
                    data["capabilities"].append(cap)
                
                # Also extract section-level offerings (H4 headings under capabilities)
                # Pattern: #### 2.1 Unified Constituent Platforms
                subsection_pattern = r"(?:^|\n)(?:####?\s*)?(\d+\.\d+)\s+([^\n]+)"
                subsections = re.findall(subsection_pattern, content)
                for num, name in subsections:
                    cap_id += 1
                    # Get description from following paragraph
                    desc_match = re.search(
                        rf"{re.escape(num)}\s+{re.escape(name)}[:\s]*\*?\*?([^\n]+(?:\n(?![-•*#\d])[^\n]+)*)",
                        content
                    )
                    desc = desc_match.group(1).strip() if desc_match else ""
                    
                    cap = {
                        "id": f"CAP-{cap_id:03d}",
                        "name": name.strip()[:100],
                        "description": desc[:500],
                        "category": section["title"],
                        "section_number": num,
                    }
                    data["offerings"].append(cap)
            
            # Look for differentiator sections
            elif any(kw in title_lower for kw in ["differentiator", "advantage", "why choose", "difference", "strategic"]):
                diff_id = len(data["differentiators"])
                
                # Find index of this section
                section_idx = sections.index(section) if section in sections else -1
                
                # Look at following subsections until we hit a same-level or higher section
                current_level = section.get("level", 0)
                if section_idx >= 0:
                    for i in range(section_idx + 1, len(sections)):
                        next_section = sections[i]
                        next_level = next_section.get("level", 0)
                        
                        # Stop if we hit a same-level or higher section
                        if next_level <= current_level:
                            break
                        
                        # This is a subsection - treat it as a differentiator
                        diff_id += 1
                        diff = {
                            "id": f"DIFF-{diff_id:03d}",
                            "title": next_section.get("title", "").strip()[:100],
                            "description": next_section.get("content", "")[:500],
                        }
                        data["differentiators"].append(diff)
                
                # Also extract bullet points from current section
                bold_bullets = re.findall(
                    r"[-•*]\s+(?:>\s*)?\*\*([^*]+)\*\*[:\s]*([^\n]*)",
                    content
                )
                for name, desc in bold_bullets:
                    diff_id += 1
                    diff = {
                        "id": f"DIFF-{diff_id:03d}",
                        "title": name.strip()[:100],
                        "description": desc.strip()[:500] if desc else "",
                    }
                    data["differentiators"].append(diff)
            
            # Look for use cases - also in the main content
            if "use case" in title_lower or "application" in title_lower:
                cases = re.findall(r"[-•*]\s+([^\n]+)", content)
                data["use_cases"].extend([c.strip() for c in cases])
            
            # Extract use cases from bullet points mentioning specific applications
            use_case_bullets = re.findall(
                r'"([^"]+)"[:\s]+([^.\n]+)',
                content
            )
            for name, desc in use_case_bullets:
                if len(name) > 10 and len(name) < 100:
                    data["use_cases"].append(f"{name}: {desc.strip()}")
            
            # Look for partnerships
            if "partner" in title_lower or "alliance" in title_lower:
                # Extract partner names
                partners = re.findall(r"(?:with|partner(?:ship)?[:\s]+)([A-Z][A-Za-z\s]+)", content)
                data["partnerships"].extend([p.strip() for p in partners if len(p.strip()) > 3])
                
                # Also look for specific named partners
                named_partners = re.findall(r"(Google|AWS|Azure|Microsoft|Salesforce|SAP|Oracle)", content)
                data["partnerships"].extend(list(set(named_partners)))
        
        # Deduplicate partnerships
        data["partnerships"] = list(set(data["partnerships"]))
        
        # Extract keywords from entire text
        data["keywords"] = self._extract_keywords(text)
        
        return data
    
    def _extract_past_performance_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract past performance data"""
        data = {
            "projects": [],
            "clients": [],
            "contract_types": [],
            "keywords": [],
        }
        
        project_id = 0
        for section in sections:
            content = section.get("content", "")
            title = section.get("title", "")
            
            # Look for project descriptions
            if any(kw in title.lower() for kw in ["project", "client", "case", "experience", "contract"]):
                project_id += 1
                project = {
                    "id": f"PP-{project_id:03d}",
                    "name": title,
                    "description": content[:1000],
                    "client": "",
                    "outcomes": [],
                }
                
                # Try to extract client name
                client_match = re.search(r"[Cc]lient[:\s]+([^\n]+)", content)
                if client_match:
                    project["client"] = client_match.group(1).strip()
                
                # Extract outcomes/results
                outcomes = re.findall(r"[-•*]\s+([^\n]+(?:result|outcome|achieve|deliver|improv)[^\n]*)", content, re.I)
                project["outcomes"] = [o.strip() for o in outcomes]
                
                data["projects"].append(project)
        
        data["keywords"] = self._extract_keywords(text)
        return data
    
    def _extract_resume_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract resume/CV data"""
        data = {
            "name": "",
            "title": "",
            "summary": "",
            "experience": [],
            "education": [],
            "certifications": [],
            "skills": [],
        }
        
        # Extract name from first lines
        lines = text.split("\n")
        for line in lines[:3]:
            line = line.strip()
            if line and not re.match(r"^#", line) and len(line.split()) <= 4:
                data["name"] = line
                break
        
        for section in sections:
            title_lower = section.get("title", "").lower()
            content = section.get("content", "")
            
            if "summary" in title_lower or "profile" in title_lower or "objective" in title_lower:
                data["summary"] = content[:500]
            
            elif "experience" in title_lower or "employment" in title_lower or "work history" in title_lower:
                # Parse experience entries
                exp_entries = re.split(r"\n(?=\*\*|(?:\d{4}|\d{2}/\d{4}))", content)
                for entry in exp_entries[:10]:
                    if entry.strip():
                        data["experience"].append(entry.strip()[:500])
            
            elif "education" in title_lower:
                edu_entries = re.findall(r"[-•*]\s*([^\n]+)", content)
                data["education"] = [e.strip() for e in edu_entries]
            
            elif "certification" in title_lower or "credential" in title_lower:
                cert_entries = re.findall(r"[-•*]\s*([^\n]+)", content)
                data["certifications"] = [c.strip() for c in cert_entries]
            
            elif "skill" in title_lower:
                skill_entries = re.findall(r"[-•*,]\s*([^,\n•*-]+)", content)
                data["skills"] = [s.strip() for s in skill_entries if len(s.strip()) > 2]
        
        return data
    
    def _extract_technical_approach_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract technical approach data"""
        data = {
            "methodology": "",
            "phases": [],
            "tools": [],
            "deliverables": [],
            "keywords": [],
        }
        
        for section in sections:
            title_lower = section.get("title", "").lower()
            content = section.get("content", "")
            
            if "methodology" in title_lower or "approach" in title_lower:
                data["methodology"] = content[:1000]
            
            elif "phase" in title_lower or "stage" in title_lower:
                phases = re.findall(r"[-•*]\s*([^\n]+)", content)
                data["phases"].extend([p.strip() for p in phases])
            
            elif "tool" in title_lower or "technolog" in title_lower:
                tools = re.findall(r"[-•*,]\s*([^,\n•*-]+)", content)
                data["tools"].extend([t.strip() for t in tools if len(t.strip()) > 2])
            
            elif "deliverable" in title_lower:
                delivs = re.findall(r"[-•*]\s*([^\n]+)", content)
                data["deliverables"].extend([d.strip() for d in delivs])
        
        data["keywords"] = self._extract_keywords(text)
        return data
    
    def _extract_corporate_info_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract corporate information"""
        data = {
            "duns": "",
            "cage_code": "",
            "naics_codes": [],
            "certifications": [],
            "contract_vehicles": [],
        }
        
        text_upper = text.upper()
        
        # DUNS number
        duns_match = re.search(r"DUNS[:\s#]*(\d{9})", text_upper)
        if duns_match:
            data["duns"] = duns_match.group(1)
        
        # CAGE code
        cage_match = re.search(r"CAGE[:\s]*([A-Z0-9]{5})", text_upper)
        if cage_match:
            data["cage_code"] = cage_match.group(1)
        
        # NAICS codes
        naics_matches = re.findall(r"(\d{6})\s*[-–]?\s*([A-Za-z][^\n,;]+)?", text)
        data["naics_codes"] = [{"code": m[0], "description": m[1].strip() if m[1] else ""} for m in naics_matches]
        
        # Certifications (8(a), HUBZone, SDVOSB, etc.)
        cert_patterns = [
            r"8\(a\)",
            r"HUBZone",
            r"SDVOSB",
            r"WOSB",
            r"EDWOSB",
            r"SDB",
            r"small\s+business",
            r"ISO\s*\d+",
            r"CMMI",
            r"FedRAMP",
        ]
        for pattern in cert_patterns:
            if re.search(pattern, text, re.I):
                data["certifications"].append(pattern.replace("\\s*", " ").replace("\\s+", " "))
        
        # Contract vehicles (GSA, SEWP, etc.)
        vehicle_patterns = [
            r"GSA\s+Schedule",
            r"GSA\s+MAS",
            r"SEWP",
            r"CIO-SP\d*",
            r"OASIS",
            r"STARS",
        ]
        for pattern in vehicle_patterns:
            if re.search(pattern, text, re.I):
                data["contract_vehicles"].append(pattern.replace("\\s+", " "))
        
        return data
    
    def _extract_generic_data(self, text: str, sections: List[Dict]) -> Dict:
        """Extract generic data from unknown document type"""
        return {
            "sections": [{"title": s["title"], "content": s.get("content", "")[:500]} for s in sections],
            "keywords": self._extract_keywords(text),
        }
    
    def _extract_keywords(self, text: str, max_keywords: int = 50) -> List[str]:
        """Extract keywords from text"""
        # Remove markdown formatting
        clean_text = re.sub(r"[#*_\[\](){}]", "", text)
        
        # Find potential keywords (capitalized phrases, technical terms)
        words = re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", text)
        
        # Count frequencies
        freq = {}
        for word in words:
            if len(word) > 3:
                freq[word] = freq.get(word, 0) + 1
        
        # Sort by frequency and return top keywords
        sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
        return [w[0] for w in sorted_words[:max_keywords]]


class CompanyLibrary:
    """
    Manages a library of parsed company documents
    """
    
    def __init__(self, storage_dir: Optional[str] = None):
        self.storage_dir = Path(storage_dir or tempfile.gettempdir()) / "propelai_company_library"
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.parser = CompanyLibraryParser()
        self.documents: Dict[str, ParsedDocument] = {}
        self.profile = CompanyProfile()
        self._load_library()
    
    def _load_library(self):
        """Load existing library from storage"""
        index_file = self.storage_dir / "library_index.json"
        if index_file.exists():
            try:
                with open(index_file, "r") as f:
                    data = json.load(f)
                    # Restore documents
                    for doc_data in data.get("documents", []):
                        doc = ParsedDocument(
                            id=doc_data["id"],
                            filename=doc_data["filename"],
                            document_type=DocumentType(doc_data["document_type"]),
                            title=doc_data["title"],
                            raw_text="",  # Don't store raw text in index
                            sections=doc_data.get("sections", []),
                            extracted_data=doc_data.get("extracted_data", {}),
                            parse_date=doc_data.get("parse_date", ""),
                        )
                        self.documents[doc.id] = doc
                    
                    # Restore profile
                    profile_data = data.get("profile", {})
                    self.profile.company_name = profile_data.get("company_name", "")
                    self.profile.executive_summary = profile_data.get("executive_summary", "")
            except Exception as e:
                print(f"Error loading library: {e}")
    
    def _save_library(self):
        """Save library to storage"""
        index_file = self.storage_dir / "library_index.json"
        data = {
            "documents": [doc.to_dict() for doc in self.documents.values()],
            "profile": self.profile.to_dict(),
        }
        with open(index_file, "w") as f:
            json.dump(data, f, indent=2)
    
    def add_document(self, file_path: str, document_type: Optional[DocumentType] = None) -> ParsedDocument:
        """
        Add a document to the library
        
        Args:
            file_path: Path to document file
            document_type: Optional type hint
            
        Returns:
            Parsed document
        """
        # Parse document
        doc = self.parser.parse_document(file_path, document_type)
        
        # Store document
        self.documents[doc.id] = doc
        
        # Copy file to storage
        dest_path = self.storage_dir / f"{doc.id}_{doc.filename}"
        import shutil
        shutil.copy2(file_path, dest_path)
        
        # Update profile
        self._update_profile(doc)
        
        # Save library
        self._save_library()
        
        return doc
    
    def _update_profile(self, doc: ParsedDocument):
        """Update company profile with data from new document"""
        data = doc.extracted_data
        
        if doc.document_type == DocumentType.CAPABILITIES:
            if data.get("company_name"):
                self.profile.company_name = data["company_name"]
            if data.get("executive_summary"):
                self.profile.executive_summary = data["executive_summary"]
            
            # Add capabilities
            for cap_data in data.get("capabilities", []):
                cap = Capability(
                    id=cap_data["id"],
                    name=cap_data["name"],
                    description=cap_data.get("description", ""),
                    category=cap_data.get("category", ""),
                )
                self.profile.capabilities.append(cap)
            
            # Add differentiators
            for diff_data in data.get("differentiators", []):
                diff = Differentiator(
                    id=diff_data["id"],
                    title=diff_data["title"],
                    description=diff_data.get("description", ""),
                )
                self.profile.differentiators.append(diff)
        
        elif doc.document_type == DocumentType.PAST_PERFORMANCE:
            for proj_data in data.get("projects", []):
                pp = PastPerformance(
                    id=proj_data["id"],
                    project_name=proj_data["name"],
                    client=proj_data.get("client", ""),
                    description=proj_data.get("description", ""),
                    outcomes=proj_data.get("outcomes", []),
                )
                self.profile.past_performance.append(pp)
        
        elif doc.document_type == DocumentType.RESUME:
            kp = KeyPersonnel(
                id=f"KP-{len(self.profile.key_personnel)+1:03d}",
                name=data.get("name", ""),
                title=data.get("title", ""),
                summary=data.get("summary", ""),
                education=data.get("education", []),
                certifications=data.get("certifications", []),
                skills=data.get("skills", []),
            )
            self.profile.key_personnel.append(kp)
        
        elif doc.document_type == DocumentType.CORPORATE_INFO:
            if data.get("naics_codes"):
                self.profile.naics_codes = [n["code"] for n in data["naics_codes"]]
            if data.get("certifications"):
                self.profile.certifications.extend(data["certifications"])
            if data.get("contract_vehicles"):
                self.profile.contract_vehicles.extend(data["contract_vehicles"])
        
        # Merge keywords
        self.profile.keywords = list(set(self.profile.keywords + data.get("keywords", [])))
    
    def get_document(self, doc_id: str) -> Optional[ParsedDocument]:
        """Get document by ID"""
        return self.documents.get(doc_id)
    
    def list_documents(self) -> List[Dict]:
        """List all documents in library"""
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "type": doc.document_type.value,
                "title": doc.title,
                "parse_date": doc.parse_date,
            }
            for doc in self.documents.values()
        ]
    
    def remove_document(self, doc_id: str) -> bool:
        """Remove document from library"""
        if doc_id in self.documents:
            del self.documents[doc_id]
            self._save_library()
            return True
        return False
    
    def get_profile(self) -> Dict:
        """Get aggregated company profile"""
        return self.profile.to_dict()
    
    def search(self, query: str) -> List[Dict]:
        """
        Search library for relevant content
        
        Args:
            query: Search query
            
        Returns:
            List of matching content
        """
        results = []
        query_lower = query.lower()
        query_terms = query_lower.split()
        
        # Search capabilities
        for cap in self.profile.capabilities:
            score = sum(1 for term in query_terms if term in cap.name.lower() or term in cap.description.lower())
            if score > 0:
                results.append({
                    "type": "capability",
                    "score": score,
                    "content": cap.to_dict(),
                })
        
        # Search past performance
        for pp in self.profile.past_performance:
            score = sum(1 for term in query_terms if term in pp.project_name.lower() or term in pp.description.lower())
            if score > 0:
                results.append({
                    "type": "past_performance",
                    "score": score,
                    "content": pp.to_dict(),
                })
        
        # Search key personnel
        for kp in self.profile.key_personnel:
            score = sum(1 for term in query_terms if term in kp.name.lower() or term in kp.summary.lower() or any(term in s.lower() for s in kp.skills))
            if score > 0:
                results.append({
                    "type": "key_personnel",
                    "score": score,
                    "content": kp.to_dict(),
                })
        
        # Search differentiators
        for diff in self.profile.differentiators:
            score = sum(1 for term in query_terms if term in diff.title.lower() or term in diff.description.lower())
            if score > 0:
                results.append({
                    "type": "differentiator",
                    "score": score,
                    "content": diff.to_dict(),
                })
        
        # Sort by score
        results.sort(key=lambda x: x["score"], reverse=True)
        
        return results


# Export for API
__all__ = [
    "DocumentType",
    "CompanyLibraryParser",
    "CompanyLibrary",
    "ParsedDocument",
    "CompanyProfile",
    "Capability",
    "PastPerformance",
    "KeyPersonnel",
    "Differentiator",
]
