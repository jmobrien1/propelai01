"""
PropelAI Cycle 5: Multi-Format Document Parser
Parses PDF, DOCX, XLSX with section boundary detection

Preserves page numbers and section references for traceability

v4.0 Phase 1: Added CoordinateExtractor for source traceability (Trust Gate)
"""

import os
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass

from .models import ParsedDocument, DocumentType, RFPBundle
from .document_structure import BoundingBox, SourceCoordinate


class CoordinateExtractor:
    """
    Extracts character-level bounding boxes from PDFs using pdfplumber.

    This is the "Trust Gate" implementation - enables visual proof of
    extraction accuracy by providing exact coordinates for any extracted text.

    Design:
    - Lazy loading: Only invokes pdfplumber when coordinates are requested
    - Caches page data to avoid re-parsing
    - Normalizes PDF coordinates to web-friendly format (0.0-1.0)
    """

    def __init__(self):
        self._page_cache: Dict[str, Dict[int, Dict]] = {}  # {filepath: {page_num: page_data}}
        self._pdfplumber_available = None

    def _check_pdfplumber(self) -> bool:
        """Check if pdfplumber is available"""
        if self._pdfplumber_available is None:
            try:
                import pdfplumber
                self._pdfplumber_available = True
            except ImportError:
                self._pdfplumber_available = False
                print("Warning: pdfplumber not installed. Coordinate extraction unavailable.")
                print("Install with: pip install pdfplumber")
        return self._pdfplumber_available

    def _get_page_data(self, filepath: str, page_num: int) -> Optional[Dict]:
        """
        Get cached page data or extract from PDF.

        Returns dict with:
        - 'chars': List of character dicts with positions
        - 'width': Page width in points
        - 'height': Page height in points
        - 'text': Full page text
        """
        if not self._check_pdfplumber():
            return None

        # Check cache
        if filepath in self._page_cache and page_num in self._page_cache[filepath]:
            return self._page_cache[filepath][page_num]

        # Extract page data
        import pdfplumber

        try:
            with pdfplumber.open(filepath) as pdf:
                if page_num < 1 or page_num > len(pdf.pages):
                    return None

                page = pdf.pages[page_num - 1]  # pdfplumber is 0-indexed

                page_data = {
                    'chars': page.chars,
                    'width': float(page.width),
                    'height': float(page.height),
                    'text': page.extract_text() or ""
                }

                # Cache the result
                if filepath not in self._page_cache:
                    self._page_cache[filepath] = {}
                self._page_cache[filepath][page_num] = page_data

                return page_data

        except Exception as e:
            print(f"Warning: Could not extract page data from {filepath}: {e}")
            return None

    def find_text_coordinates(
        self,
        filepath: str,
        search_text: str,
        page_num: Optional[int] = None,
        fuzzy_match: bool = True
    ) -> Optional[SourceCoordinate]:
        """
        Find the coordinates of a text string in a PDF.

        Args:
            filepath: Path to the PDF file
            search_text: The text to find
            page_num: Specific page to search (None = search all pages)
            fuzzy_match: If True, normalizes whitespace for matching

        Returns:
            SourceCoordinate with bounding boxes, or None if not found
        """
        if not self._check_pdfplumber():
            return None

        import pdfplumber

        # Normalize search text if fuzzy matching
        if fuzzy_match:
            search_normalized = ' '.join(search_text.split())
        else:
            search_normalized = search_text

        try:
            with pdfplumber.open(filepath) as pdf:
                pages_to_search = [page_num] if page_num else range(1, len(pdf.pages) + 1)

                for pn in pages_to_search:
                    if pn < 1 or pn > len(pdf.pages):
                        continue

                    page = pdf.pages[pn - 1]
                    page_text = page.extract_text() or ""

                    # Normalize page text if fuzzy matching
                    if fuzzy_match:
                        page_normalized = ' '.join(page_text.split())
                    else:
                        page_normalized = page_text

                    # Find the text in the page
                    idx = page_normalized.find(search_normalized)
                    if idx == -1:
                        continue

                    # Map back to original character positions
                    chars = page.chars
                    if not chars:
                        continue

                    # Build a mapping from normalized text position to char indices
                    char_positions = self._find_char_positions(
                        chars, search_normalized, page_text, fuzzy_match
                    )

                    if char_positions:
                        return SourceCoordinate.from_text_match(
                            document_id=os.path.basename(filepath),
                            page_index=pn,
                            char_positions=char_positions,
                            page_width=float(page.width),
                            page_height=float(page.height)
                        )

                return None

        except Exception as e:
            print(f"Warning: Error finding text coordinates: {e}")
            return None

    def _find_char_positions(
        self,
        chars: List[Dict],
        search_text: str,
        full_text: str,
        fuzzy_match: bool
    ) -> List[Dict[str, float]]:
        """
        Find character position dicts for matching text.

        Returns list of dicts with x0, top, x1, bottom keys.
        """
        if not chars:
            return []

        # Build text from chars
        char_text = ''.join(c.get('text', '') for c in chars)

        if fuzzy_match:
            # Map char indices to normalized positions
            normalized_to_char_idx = []
            norm_text = []
            in_whitespace = False

            for i, c in enumerate(chars):
                char = c.get('text', '')
                if char.isspace():
                    if not in_whitespace:
                        norm_text.append(' ')
                        normalized_to_char_idx.append(i)
                        in_whitespace = True
                else:
                    norm_text.append(char)
                    normalized_to_char_idx.append(i)
                    in_whitespace = False

            normalized_text = ''.join(norm_text)
            search_normalized = ' '.join(search_text.split())

            # Find in normalized text
            idx = normalized_text.find(search_normalized)
            if idx == -1:
                return []

            # Get the char indices for the match
            start_char_idx = normalized_to_char_idx[idx] if idx < len(normalized_to_char_idx) else 0
            end_idx = idx + len(search_normalized)
            end_char_idx = normalized_to_char_idx[min(end_idx - 1, len(normalized_to_char_idx) - 1)] + 1

            # Extract the matching chars
            matching_chars = chars[start_char_idx:end_char_idx]
        else:
            # Direct match
            idx = char_text.find(search_text)
            if idx == -1:
                return []
            matching_chars = chars[idx:idx + len(search_text)]

        # Convert to position dicts
        positions = []
        for c in matching_chars:
            if c.get('text', '').strip():  # Skip pure whitespace chars
                positions.append({
                    'x0': c.get('x0', 0),
                    'top': c.get('top', 0),
                    'x1': c.get('x1', c.get('x0', 0) + 5),
                    'bottom': c.get('bottom', c.get('top', 0) + 10)
                })

        return positions

    def get_page_dimensions(self, filepath: str, page_num: int) -> Optional[Tuple[float, float]]:
        """Get page dimensions (width, height) in points."""
        page_data = self._get_page_data(filepath, page_num)
        if page_data:
            return (page_data['width'], page_data['height'])
        return None

    def clear_cache(self, filepath: Optional[str] = None):
        """Clear cached page data."""
        if filepath:
            self._page_cache.pop(filepath, None)
        else:
            self._page_cache.clear()


# Global coordinator extractor instance
_coordinate_extractor: Optional[CoordinateExtractor] = None


def get_coordinate_extractor() -> CoordinateExtractor:
    """Get or create the global CoordinateExtractor instance."""
    global _coordinate_extractor
    if _coordinate_extractor is None:
        _coordinate_extractor = CoordinateExtractor()
    return _coordinate_extractor


class MultiFormatParser:
    """
    Parse PDF, DOCX, XLSX into unified format with section detection
    
    Uses:
    - pypdf for PDFs
    - python-docx for DOCX
    - openpyxl for XLSX
    - markitdown as fallback for complex layouts
    """
    
    # Section header patterns (federal RFP standard)
    SECTION_PATTERNS = {
        # FAR standard sections
        "section_a": r"SECTION\s*A[\s:\-–]+|SOLICITATION.*FORM",
        "section_b": r"SECTION\s*B[\s:\-–]+|SUPPLIES\s*OR\s*SERVICES",
        "section_c": r"SECTION\s*C[\s:\-–]+|DESCRIPTION.*SPECIFICATIONS|STATEMENT\s*OF\s*WORK",
        "section_d": r"SECTION\s*D[\s:\-–]+|PACKAGING.*MARKING",
        "section_e": r"SECTION\s*E[\s:\-–]+|INSPECTION.*ACCEPTANCE",
        "section_f": r"SECTION\s*F[\s:\-–]+|DELIVERIES.*PERFORMANCE",
        "section_g": r"SECTION\s*G[\s:\-–]+|CONTRACT\s*ADMINISTRATION",
        "section_h": r"SECTION\s*H[\s:\-–]+|SPECIAL\s*CONTRACT",
        "section_i": r"SECTION\s*I[\s:\-–]+|CONTRACT\s*CLAUSES",
        "section_j": r"SECTION\s*J[\s:\-–]+|ATTACHMENTS",
        "section_k": r"SECTION\s*K[\s:\-–]+|REPRESENTATIONS",
        "section_l": r"SECTION\s*L[\s:\-–]+|INSTRUCTIONS.*OFFERORS",
        "section_m": r"SECTION\s*M[\s:\-–]+|EVALUATION\s*FACTORS",
    }
    
    # Article patterns within sections
    ARTICLE_PATTERN = r"ARTICLE\s+([A-Z]\.\d+(?:\.\d+)?)"
    
    # Subsection patterns: C.3.1, L.4.a.2, etc.
    SUBSECTION_PATTERN = r"([A-Z])\.(\d+)(?:\.(\d+|[a-z]))?(?:\.(\d+|[a-z]))?"
    
    def __init__(self, use_markitdown_fallback: bool = True):
        self.use_markitdown_fallback = use_markitdown_fallback
        self.stats = {
            "pdf_parsed": 0,
            "docx_parsed": 0,
            "xlsx_parsed": 0,
            "parse_failures": 0,
        }
    
    def parse_bundle(self, bundle: RFPBundle) -> Dict[str, ParsedDocument]:
        """
        Parse all documents in the bundle
        
        Returns:
            Dict mapping document role to ParsedDocument
        """
        results = {}
        
        # Parse main solicitation
        if bundle.main_document:
            parsed = self.parse_file(bundle.main_document, DocumentType.MAIN_SOLICITATION)
            if parsed:
                results["main"] = parsed
                bundle.parsed_documents["main"] = parsed
        
        # Parse SOW if separate
        if bundle.sow_document:
            parsed = self.parse_file(bundle.sow_document, DocumentType.STATEMENT_OF_WORK)
            if parsed:
                results["sow"] = parsed
                bundle.parsed_documents["sow"] = parsed
        
        # Parse amendments in order
        for i, amendment_path in enumerate(bundle.amendments):
            parsed = self.parse_file(amendment_path, DocumentType.AMENDMENT)
            if parsed:
                key = f"amendment_{i+1}"
                results[key] = parsed
                bundle.parsed_documents[key] = parsed
        
        # Parse research outlines (NIH-specific)
        for ro_id, ro_path in bundle.research_outlines.items():
            parsed = self.parse_file(ro_path, DocumentType.RESEARCH_OUTLINE)
            if parsed:
                results[ro_id] = parsed
                bundle.parsed_documents[ro_id] = parsed
        
        # Parse attachments
        for att_id, att_path in bundle.attachments.items():
            parsed = self.parse_file(att_path, DocumentType.ATTACHMENT)
            if parsed:
                results[att_id] = parsed
                bundle.parsed_documents[att_id] = parsed
        
        return results
    
    def parse_file(self, filepath: str, doc_type: DocumentType) -> Optional[ParsedDocument]:
        """
        Parse a single file based on extension
        
        Returns:
            ParsedDocument or None if parsing failed
        """
        if not os.path.exists(filepath):
            return None
        
        ext = os.path.splitext(filepath)[1].lower()
        filename = os.path.basename(filepath)
        
        try:
            if ext == ".pdf":
                return self._parse_pdf(filepath, filename, doc_type)
            elif ext in [".docx", ".doc"]:
                return self._parse_docx(filepath, filename, doc_type)
            elif ext in [".xlsx", ".xls"]:
                return self._parse_xlsx(filepath, filename, doc_type)
            elif ext == ".txt":
                return self._parse_text(filepath, filename, doc_type)
            else:
                # Try markitdown for unknown formats
                if self.use_markitdown_fallback:
                    return self._parse_with_markitdown(filepath, filename, doc_type)
                return None
        except Exception as e:
            self.stats["parse_failures"] += 1
            print(f"Warning: Failed to parse {filepath}: {e}")
            return None
    
    def _parse_pdf(self, filepath: str, filename: str, doc_type: DocumentType) -> ParsedDocument:
        """Parse PDF using pypdf"""
        from pypdf import PdfReader
        
        reader = PdfReader(filepath)
        pages = []
        full_text_parts = []
        
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
            full_text_parts.append(text)
        
        full_text = "\n\n".join(full_text_parts)
        
        # Detect sections
        sections = self._detect_sections(full_text)
        
        # Extract tables (basic detection)
        tables = self._extract_table_hints(full_text)
        
        # Calculate extraction quality
        quality = self._assess_extraction_quality(full_text, len(reader.pages))
        
        self.stats["pdf_parsed"] += 1
        
        return ParsedDocument(
            filepath=filepath,
            filename=filename,
            document_type=doc_type,
            full_text=full_text,
            pages=pages,
            page_count=len(pages),
            sections=sections,
            tables=tables,
            title=self._extract_title(full_text),
            parser_used="pypdf",
            extraction_quality=quality,
        )
    
    def _parse_docx(self, filepath: str, filename: str, doc_type: DocumentType) -> ParsedDocument:
        """Parse DOCX using python-docx"""
        from docx import Document
        
        doc = Document(filepath)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        full_text = "\n\n".join(paragraphs)
        
        # For DOCX, we don't have natural page breaks, estimate
        pages = self._split_into_pages(full_text, chars_per_page=3000)
        
        sections = self._detect_sections(full_text)
        tables = self._extract_docx_tables(doc)
        
        self.stats["docx_parsed"] += 1
        
        return ParsedDocument(
            filepath=filepath,
            filename=filename,
            document_type=doc_type,
            full_text=full_text,
            pages=pages,
            page_count=len(pages),
            sections=sections,
            tables=tables,
            title=self._extract_title(full_text),
            parser_used="python-docx",
            extraction_quality=1.0,
        )
    
    def _parse_xlsx(self, filepath: str, filename: str, doc_type: DocumentType) -> ParsedDocument:
        """Parse Excel using openpyxl"""
        from openpyxl import load_workbook
        
        wb = load_workbook(filepath, data_only=True)
        
        tables = []
        full_text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            
            # Extract as table
            table_data = {
                "sheet_name": sheet_name,
                "headers": [],
                "rows": [],
            }
            
            rows_text = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                
                # Skip completely empty rows
                if not any(row_values):
                    continue
                
                if i == 0:
                    table_data["headers"] = row_values
                else:
                    table_data["rows"].append(row_values)
                
                rows_text.append(" | ".join(row_values))
            
            tables.append(table_data)
            full_text_parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))
        
        full_text = "\n\n".join(full_text_parts)
        
        self.stats["xlsx_parsed"] += 1
        
        return ParsedDocument(
            filepath=filepath,
            filename=filename,
            document_type=doc_type,
            full_text=full_text,
            pages=[full_text],  # Excel doesn't have pages
            page_count=1,
            sections={},
            tables=tables,
            title=filename,
            parser_used="openpyxl",
            extraction_quality=1.0,
        )
    
    def _parse_text(self, filepath: str, filename: str, doc_type: DocumentType) -> ParsedDocument:
        """Parse plain text file"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()
        
        pages = self._split_into_pages(full_text)
        sections = self._detect_sections(full_text)
        
        return ParsedDocument(
            filepath=filepath,
            filename=filename,
            document_type=doc_type,
            full_text=full_text,
            pages=pages,
            page_count=len(pages),
            sections=sections,
            tables=[],
            parser_used="text",
            extraction_quality=1.0,
        )
    
    def _parse_with_markitdown(self, filepath: str, filename: str, doc_type: DocumentType) -> Optional[ParsedDocument]:
        """Fallback parser using markitdown"""
        try:
            from markitdown import MarkItDown
            
            md = MarkItDown()
            result = md.convert(filepath)
            
            full_text = result.text_content
            pages = self._split_into_pages(full_text)
            sections = self._detect_sections(full_text)
            
            return ParsedDocument(
                filepath=filepath,
                filename=filename,
                document_type=doc_type,
                full_text=full_text,
                pages=pages,
                page_count=len(pages),
                sections=sections,
                tables=[],
                parser_used="markitdown",
                extraction_quality=0.9,
            )
        except Exception as e:
            print(f"Warning: markitdown failed for {filepath}: {e}")
            return None
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """
        Detect section boundaries in the document
        
        Returns:
            Dict mapping section ID to section text
        """
        sections = {}
        
        # Find all section headers and their positions
        section_positions = []
        
        for section_id, pattern in self.SECTION_PATTERNS.items():
            for match in re.finditer(pattern, text, re.IGNORECASE):
                section_positions.append((match.start(), section_id, match.group()))
        
        # Sort by position
        section_positions.sort(key=lambda x: x[0])
        
        # Extract text between sections
        for i, (start_pos, section_id, header) in enumerate(section_positions):
            # Find end position (start of next section or end of doc)
            if i + 1 < len(section_positions):
                end_pos = section_positions[i + 1][0]
            else:
                end_pos = len(text)
            
            section_text = text[start_pos:end_pos]
            sections[section_id] = section_text
        
        return sections
    
    def _extract_table_hints(self, text: str) -> List[Dict[str, Any]]:
        """
        Detect potential tables in PDF text
        
        This is heuristic-based; for accurate table extraction,
        would need a table detection library
        """
        tables = []
        
        # Look for patterns that suggest tabular data
        # Multiple columns separated by whitespace
        lines = text.split('\n')
        
        table_start = None
        table_lines = []
        
        for i, line in enumerate(lines):
            # Heuristic: line with 3+ segments of 2+ spaces
            segments = re.split(r'\s{2,}', line.strip())
            if len(segments) >= 3 and all(len(s) > 0 for s in segments[:3]):
                if table_start is None:
                    table_start = i
                table_lines.append(segments)
            else:
                if table_start is not None and len(table_lines) >= 2:
                    tables.append({
                        "start_line": table_start,
                        "end_line": i,
                        "rows": table_lines,
                        "headers": table_lines[0] if table_lines else [],
                    })
                table_start = None
                table_lines = []
        
        return tables
    
    def _extract_docx_tables(self, doc) -> List[Dict[str, Any]]:
        """Extract tables from DOCX document"""
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "table_index": i,
                "headers": [],
                "rows": [],
            }
            
            for j, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if j == 0:
                    table_data["headers"] = row_data
                else:
                    table_data["rows"].append(row_data)
            
            tables.append(table_data)
        
        return tables
    
    def _split_into_pages(self, text: str, chars_per_page: int = 3000) -> List[str]:
        """Split text into estimated pages"""
        pages = []
        for i in range(0, len(text), chars_per_page):
            pages.append(text[i:i + chars_per_page])
        return pages if pages else [text]
    
    def _extract_title(self, text: str) -> Optional[str]:
        """Extract document title from first few lines"""
        lines = text.split('\n')[:20]
        
        for line in lines:
            line = line.strip()
            # Look for title-like patterns
            if len(line) > 10 and len(line) < 200:
                if re.search(r"REQUEST\s*FOR\s*PROPOSAL|RFP|SOLICITATION|STATEMENT\s*OF\s*WORK", line, re.IGNORECASE):
                    return line
        
        # Return first substantial line
        for line in lines:
            if len(line.strip()) > 20:
                return line.strip()[:100]
        
        return None
    
    def _assess_extraction_quality(self, text: str, page_count: int) -> float:
        """
        Assess quality of PDF text extraction
        
        Returns:
            Float 0-1, where 1 is high quality
        """
        if not text:
            return 0.0
        
        # Check for garbled text indicators
        garbled_patterns = [
            r'[^\x00-\x7F]{5,}',  # Long non-ASCII sequences
            r'\s{10,}',          # Excessive whitespace
            r'([a-z])\1{4,}',    # Repeated characters
        ]
        
        quality = 1.0
        
        for pattern in garbled_patterns:
            matches = len(re.findall(pattern, text[:5000]))
            if matches > 10:
                quality -= 0.2
        
        # Check average chars per page
        chars_per_page = len(text) / max(page_count, 1)
        if chars_per_page < 500:  # Likely OCR issues
            quality -= 0.3
        
        return max(0.0, quality)
    
    def find_page_for_text(self, doc: ParsedDocument, search_text: str) -> int:
        """Find which page contains the given text"""
        search_lower = search_text.lower()[:100]  # First 100 chars
        
        for i, page_text in enumerate(doc.pages):
            if search_lower in page_text.lower():
                return i + 1  # 1-indexed
        
        return 0  # Not found
    
    def get_section_for_position(self, doc: ParsedDocument, position: int) -> str:
        """Determine which section contains the given character position"""
        current_pos = 0
        
        for section_id, section_text in doc.sections.items():
            section_start = doc.full_text.find(section_text)
            if section_start <= position < section_start + len(section_text):
                return section_id.replace("section_", "").upper()
        
        return "UNKNOWN"
