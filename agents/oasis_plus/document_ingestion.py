"""
Document Ingestion Pipeline with OCR
====================================

Multi-format document ingestion with OCR support for scanned PDFs.

Supports:
- Native PDF text extraction (PyMuPDF)
- OCR for scanned documents (Tesseract)
- DOCX parsing (python-docx)
- Excel parsing (openpyxl)
- Automatic quality assessment
- Chunking with position tracking

The dual-engine approach uses PyMuPDF for native text and
falls back to Tesseract OCR for scanned pages.
"""

import logging
import os
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Dict, Optional, Any, Tuple, Generator
from datetime import datetime

logger = logging.getLogger(__name__)

# Feature availability flags
PYMUPDF_AVAILABLE = False
TESSERACT_AVAILABLE = False
DOCX_AVAILABLE = False
OPENPYXL_AVAILABLE = False
PDF2IMAGE_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    pass

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    pass

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    pass

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    pass

from .models import DocumentChunk


@dataclass
class PageContent:
    """Content extracted from a single page"""
    page_number: int
    text: str
    bboxes: List[Tuple[str, Tuple[float, float, float, float]]] = field(default_factory=list)
    was_ocr: bool = False
    ocr_confidence: Optional[float] = None
    image_path: Optional[str] = None


@dataclass
class DocumentContent:
    """Complete content from a document"""
    filename: str
    filepath: str
    pages: List[PageContent] = field(default_factory=list)
    full_text: str = ""
    page_count: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_method: str = "unknown"
    quality_score: float = 1.0
    processing_time_ms: int = 0


@dataclass
class IngestionResult:
    """Result of document ingestion"""
    document: DocumentContent
    chunks: List[DocumentChunk] = field(default_factory=list)
    success: bool = True
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


class DocumentIngestionPipeline:
    """
    Multi-format document ingestion with OCR support.

    Uses a dual-engine approach:
    1. PyMuPDF for native PDF text extraction
    2. Tesseract OCR for scanned pages

    Automatically detects page quality and uses OCR when needed.
    """

    def __init__(
        self,
        use_ocr: bool = True,
        ocr_threshold: float = 0.3,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        tesseract_config: str = "--oem 3 --psm 6",
        embedding_function: Optional[callable] = None,
    ):
        """
        Initialize the ingestion pipeline.

        Args:
            use_ocr: Whether to use OCR for low-quality pages
            ocr_threshold: Text density threshold below which OCR is used
            chunk_size: Characters per chunk
            chunk_overlap: Overlap between chunks
            tesseract_config: Tesseract configuration string
            embedding_function: Optional function to generate embeddings
        """
        self.use_ocr = use_ocr and TESSERACT_AVAILABLE
        self.ocr_threshold = ocr_threshold
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tesseract_config = tesseract_config
        self.embedding_function = embedding_function

        # Verify OCR availability
        if use_ocr and not TESSERACT_AVAILABLE:
            logger.warning("OCR requested but Tesseract not available")

        if use_ocr and not PDF2IMAGE_AVAILABLE:
            logger.warning("pdf2image not available - OCR may be limited")

    def ingest(
        self,
        filepath: str,
        project_id: str,
        document_id: Optional[str] = None,
    ) -> IngestionResult:
        """
        Ingest a document and create chunks.

        Args:
            filepath: Path to the document file
            project_id: Project ID for chunk association
            document_id: Optional document ID (defaults to filename)

        Returns:
            IngestionResult with document content and chunks
        """
        start_time = datetime.now()
        path = Path(filepath)

        if not path.exists():
            return IngestionResult(
                document=DocumentContent(
                    filename=path.name,
                    filepath=filepath,
                ),
                success=False,
                errors=[f"File not found: {filepath}"],
            )

        document_id = document_id or path.name
        extension = path.suffix.lower()

        # Route to appropriate parser
        if extension == '.pdf':
            result = self._ingest_pdf(filepath, project_id, document_id)
        elif extension in ['.docx', '.doc']:
            result = self._ingest_docx(filepath, project_id, document_id)
        elif extension in ['.xlsx', '.xls']:
            result = self._ingest_excel(filepath, project_id, document_id)
        else:
            result = IngestionResult(
                document=DocumentContent(
                    filename=path.name,
                    filepath=filepath,
                ),
                success=False,
                errors=[f"Unsupported file format: {extension}"],
            )

        # Calculate processing time
        duration = datetime.now() - start_time
        result.document.processing_time_ms = int(duration.total_seconds() * 1000)

        return result

    def _ingest_pdf(
        self,
        filepath: str,
        project_id: str,
        document_id: str,
    ) -> IngestionResult:
        """Ingest a PDF document"""
        if not PYMUPDF_AVAILABLE:
            return IngestionResult(
                document=DocumentContent(
                    filename=Path(filepath).name,
                    filepath=filepath,
                ),
                success=False,
                errors=["PyMuPDF not available for PDF processing"],
            )

        result = IngestionResult(
            document=DocumentContent(
                filename=Path(filepath).name,
                filepath=filepath,
            )
        )

        try:
            doc = fitz.open(filepath)
            result.document.page_count = doc.page_count
            result.document.metadata = dict(doc.metadata) if doc.metadata else {}

            all_text_parts = []

            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_content = self._extract_page_content(page, page_num + 1, filepath)
                result.document.pages.append(page_content)
                all_text_parts.append(page_content.text)

            doc.close()

            # Combine all text
            result.document.full_text = "\n\n".join(all_text_parts)
            result.document.extraction_method = "pymupdf" + ("+ocr" if any(p.was_ocr for p in result.document.pages) else "")

            # Calculate quality score
            result.document.quality_score = self._calculate_quality_score(result.document)

            # Create chunks
            result.chunks = self._create_chunks(
                result.document, project_id, document_id
            )

        except Exception as e:
            result.success = False
            result.errors.append(f"PDF processing failed: {e}")
            logger.error(f"PDF ingestion error: {e}")

        return result

    def _extract_page_content(
        self,
        page: "fitz.Page",
        page_number: int,
        filepath: str,
    ) -> PageContent:
        """Extract content from a single PDF page"""
        # Try native text extraction first
        text = page.get_text()
        text_blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        # Extract bounding boxes
        bboxes = []
        if "blocks" in text_blocks:
            for block in text_blocks["blocks"]:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            span_text = span.get("text", "").strip()
                            if span_text:
                                bbox = span.get("bbox", (0, 0, 0, 0))
                                bboxes.append((span_text, tuple(bbox)))

        # Check if OCR is needed
        text_density = len(text.strip()) / max(1, page.rect.width * page.rect.height)
        needs_ocr = self.use_ocr and text_density < self.ocr_threshold

        page_content = PageContent(
            page_number=page_number,
            text=text,
            bboxes=bboxes,
            was_ocr=False,
        )

        if needs_ocr and TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE:
            logger.info(f"Page {page_number} needs OCR (density: {text_density:.4f})")
            ocr_content = self._ocr_page(page, page_number, filepath)
            if ocr_content and len(ocr_content.text) > len(text):
                page_content = ocr_content

        return page_content

    def _ocr_page(
        self,
        page: "fitz.Page",
        page_number: int,
        filepath: str,
    ) -> Optional[PageContent]:
        """Perform OCR on a page"""
        try:
            # Convert page to image
            pix = page.get_pixmap(dpi=300)
            img_path = tempfile.mktemp(suffix=".png")
            pix.save(img_path)

            # Perform OCR
            img = Image.open(img_path)
            ocr_data = pytesseract.image_to_data(
                img,
                config=self.tesseract_config,
                output_type=pytesseract.Output.DICT
            )

            # Extract text and bounding boxes
            text_parts = []
            bboxes = []
            confidences = []

            for i in range(len(ocr_data['text'])):
                word = ocr_data['text'][i].strip()
                conf = int(ocr_data['conf'][i])

                if word and conf > 0:
                    text_parts.append(word)
                    confidences.append(conf)

                    # Get bbox
                    x = ocr_data['left'][i]
                    y = ocr_data['top'][i]
                    w = ocr_data['width'][i]
                    h = ocr_data['height'][i]
                    bboxes.append((word, (x, y, x + w, y + h)))

            # Clean up
            os.remove(img_path)

            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            return PageContent(
                page_number=page_number,
                text=" ".join(text_parts),
                bboxes=bboxes,
                was_ocr=True,
                ocr_confidence=avg_confidence / 100,  # Normalize to 0-1
            )

        except Exception as e:
            logger.error(f"OCR failed for page {page_number}: {e}")
            return None

    def _ingest_docx(
        self,
        filepath: str,
        project_id: str,
        document_id: str,
    ) -> IngestionResult:
        """Ingest a Word document"""
        if not DOCX_AVAILABLE:
            return IngestionResult(
                document=DocumentContent(
                    filename=Path(filepath).name,
                    filepath=filepath,
                ),
                success=False,
                errors=["python-docx not available for DOCX processing"],
            )

        result = IngestionResult(
            document=DocumentContent(
                filename=Path(filepath).name,
                filepath=filepath,
            )
        )

        try:
            doc = Document(filepath)

            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)

            # Estimate page count (rough approximation)
            chars_per_page = 3000
            estimated_pages = max(1, len(full_text) // chars_per_page)

            # Create single "page" for DOCX
            result.document.pages = [
                PageContent(
                    page_number=i + 1,
                    text=full_text[i * chars_per_page:(i + 1) * chars_per_page],
                )
                for i in range(estimated_pages)
            ]

            result.document.full_text = full_text
            result.document.page_count = estimated_pages
            result.document.extraction_method = "python-docx"
            result.document.quality_score = 1.0  # Native extraction is high quality

            # Create chunks
            result.chunks = self._create_chunks(
                result.document, project_id, document_id
            )

        except Exception as e:
            result.success = False
            result.errors.append(f"DOCX processing failed: {e}")
            logger.error(f"DOCX ingestion error: {e}")

        return result

    def _ingest_excel(
        self,
        filepath: str,
        project_id: str,
        document_id: str,
    ) -> IngestionResult:
        """Ingest an Excel document"""
        if not OPENPYXL_AVAILABLE:
            return IngestionResult(
                document=DocumentContent(
                    filename=Path(filepath).name,
                    filepath=filepath,
                ),
                success=False,
                errors=["openpyxl not available for Excel processing"],
            )

        result = IngestionResult(
            document=DocumentContent(
                filename=Path(filepath).name,
                filepath=filepath,
            )
        )

        try:
            wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)

            all_text_parts = []
            page_num = 0

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                page_num += 1

                sheet_text_parts = [f"=== Sheet: {sheet_name} ==="]

                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(v for v in row_values if v)
                    if row_text:
                        sheet_text_parts.append(row_text)

                sheet_text = "\n".join(sheet_text_parts)
                all_text_parts.append(sheet_text)

                result.document.pages.append(PageContent(
                    page_number=page_num,
                    text=sheet_text,
                ))

            wb.close()

            result.document.full_text = "\n\n".join(all_text_parts)
            result.document.page_count = page_num
            result.document.extraction_method = "openpyxl"
            result.document.quality_score = 1.0

            # Create chunks
            result.chunks = self._create_chunks(
                result.document, project_id, document_id
            )

        except Exception as e:
            result.success = False
            result.errors.append(f"Excel processing failed: {e}")
            logger.error(f"Excel ingestion error: {e}")

        return result

    def _calculate_quality_score(self, document: DocumentContent) -> float:
        """Calculate overall document quality score"""
        if not document.pages:
            return 0.0

        scores = []
        for page in document.pages:
            if page.was_ocr:
                # OCR quality based on confidence
                scores.append(page.ocr_confidence or 0.5)
            else:
                # Native extraction assumed high quality
                text_len = len(page.text.strip())
                if text_len > 100:
                    scores.append(1.0)
                elif text_len > 0:
                    scores.append(0.7)
                else:
                    scores.append(0.0)

        return sum(scores) / len(scores) if scores else 0.0

    def _create_chunks(
        self,
        document: DocumentContent,
        project_id: str,
        document_id: str,
    ) -> List[DocumentChunk]:
        """Create document chunks with optional embeddings"""
        chunks = []

        for page in document.pages:
            text = page.text
            if not text.strip():
                continue

            # Simple chunking by character count
            start = 0
            chunk_idx = 0

            while start < len(text):
                end = min(start + self.chunk_size, len(text))

                # Try to break at sentence boundary
                if end < len(text):
                    for sep in ['. ', '.\n', '\n\n', '\n', ' ']:
                        last_sep = text.rfind(sep, start, end)
                        if last_sep > start + self.chunk_size // 2:
                            end = last_sep + len(sep)
                            break

                chunk_text = text[start:end].strip()

                if chunk_text:
                    chunk_id = f"{project_id}-{document_id}-p{page.page_number}-c{chunk_idx}"

                    # Find bbox for this chunk if available
                    chunk_bbox = None
                    if page.bboxes:
                        # Use bbox of first matching text
                        for bbox_text, bbox in page.bboxes:
                            if bbox_text in chunk_text[:100]:
                                chunk_bbox = bbox
                                break

                    # Generate embedding if function available
                    embedding = None
                    if self.embedding_function:
                        try:
                            embedding = self.embedding_function(chunk_text)
                        except Exception as e:
                            logger.warning(f"Embedding generation failed: {e}")

                    chunk = DocumentChunk(
                        chunk_id=chunk_id,
                        document_id=document_id,
                        project_id=project_id,
                        content=chunk_text,
                        embedding=embedding,
                        page_number=page.page_number,
                        bbox=chunk_bbox,
                        chunk_index=len(chunks),
                        char_start=start,
                        char_end=end,
                        ocr_confidence=page.ocr_confidence,
                        was_ocr=page.was_ocr,
                    )
                    chunks.append(chunk)
                    chunk_idx += 1

                start = end - self.chunk_overlap

        return chunks

    def batch_ingest(
        self,
        filepaths: List[str],
        project_id: str,
    ) -> List[IngestionResult]:
        """
        Ingest multiple documents.

        Args:
            filepaths: List of file paths to ingest
            project_id: Project ID for chunk association

        Returns:
            List of IngestionResult objects
        """
        results = []
        for filepath in filepaths:
            logger.info(f"Ingesting: {filepath}")
            result = self.ingest(filepath, project_id)
            results.append(result)

            if result.success:
                logger.info(
                    f"  - {result.document.page_count} pages, "
                    f"{len(result.chunks)} chunks, "
                    f"quality: {result.document.quality_score:.2f}"
                )
            else:
                logger.error(f"  - Failed: {result.errors}")

        return results


def ingest_document(
    filepath: str,
    project_id: str,
    use_ocr: bool = True,
    embedding_function: Optional[callable] = None,
) -> IngestionResult:
    """
    Convenience function to ingest a single document.

    Args:
        filepath: Path to the document
        project_id: Project ID for chunk association
        use_ocr: Whether to use OCR for scanned pages
        embedding_function: Optional function for generating embeddings

    Returns:
        IngestionResult with document content and chunks
    """
    pipeline = DocumentIngestionPipeline(
        use_ocr=use_ocr,
        embedding_function=embedding_function,
    )
    return pipeline.ingest(filepath, project_id)


def check_ocr_availability() -> Dict[str, bool]:
    """Check availability of OCR components"""
    return {
        "pymupdf": PYMUPDF_AVAILABLE,
        "tesseract": TESSERACT_AVAILABLE,
        "pdf2image": PDF2IMAGE_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "openpyxl": OPENPYXL_AVAILABLE,
        "ocr_ready": TESSERACT_AVAILABLE and PDF2IMAGE_AVAILABLE,
    }
