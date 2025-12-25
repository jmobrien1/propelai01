"""
PropelAI Document Parser
Unified parser for PDF, DOCX, and XLSX with traceability
"""

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Union
from pathlib import Path
from enum import Enum

from parsing.pdf_parser import PDFParser, PDFDocument, TextBlock, BoundingBox


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    UNKNOWN = "unknown"


@dataclass
class DocumentSection:
    """A section within a document."""
    id: str
    title: str
    content: str
    page_start: int
    page_end: int
    level: int = 1  # Heading level (1=top)
    children: List["DocumentSection"] = field(default_factory=list)
    bbox: Optional[Dict[str, Any]] = None


@dataclass
class ExtractedText:
    """Text extracted with source location."""
    text: str
    page: int
    bbox: Optional[Dict[str, Any]] = None
    section: Optional[str] = None
    confidence: float = 1.0
    context_before: str = ""
    context_after: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {
            "text": self.text,
            "page": self.page,
            "bbox": self.bbox,
            "section": self.section,
            "confidence": self.confidence,
            "context": {
                "before": self.context_before,
                "after": self.context_after,
            }
        }


@dataclass
class ParsedDocument:
    """Fully parsed document with traceability."""
    filename: str
    doc_type: DocumentType
    total_pages: int
    full_text: str
    sections: List[DocumentSection] = field(default_factory=list)
    extracted_texts: List[ExtractedText] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)

    # For PDFs, store the raw parsed document for bbox lookups
    _pdf_document: Optional[PDFDocument] = field(default=None, repr=False)

    def find_text_source(self, text: str) -> Optional[Dict[str, Any]]:
        """
        Find where a piece of text appears in the document.
        Returns source location with bbox for PDF verification.
        """
        if self._pdf_document:
            return self._pdf_document.find_requirement_source(text)

        # For non-PDF, search through extracted texts
        text_lower = text.lower().strip()
        for ext in self.extracted_texts:
            if text_lower in ext.text.lower():
                return {
                    "found": True,
                    "page": ext.page,
                    "bbox": ext.bbox,
                    "matched_text": ext.text,
                    "section": ext.section,
                    "confidence": ext.confidence,
                }

        return {"found": False, "confidence": 0.0}

    def get_page_text(self, page_number: int) -> Optional[str]:
        """Get text content of a specific page."""
        if self._pdf_document:
            page = self._pdf_document.get_page(page_number)
            return page.full_text if page else None

        # For non-PDF, approximate from extracted texts
        texts = [et.text for et in self.extracted_texts if et.page == page_number]
        return "\n\n".join(texts) if texts else None

    def search(self, query: str) -> List[ExtractedText]:
        """Search for text across the document."""
        query_lower = query.lower()
        results = []

        if self._pdf_document:
            for page_num, block in self._pdf_document.search_text(query):
                results.append(ExtractedText(
                    text=block.text,
                    page=page_num,
                    bbox=block.bbox.to_dict(),
                    context_before=block.context_before,
                    context_after=block.context_after,
                ))
        else:
            for ext in self.extracted_texts:
                if query_lower in ext.text.lower():
                    results.append(ext)

        return results


class DocumentParser:
    """
    Unified document parser with source traceability.
    Supports PDF, DOCX, and XLSX formats.
    """

    # FAR Section patterns for RFP structure detection
    SECTION_PATTERNS = {
        'A': r'(?:SECTION\s*A|PART\s*I)[\s\-:]*SOLICITATION',
        'B': r'(?:SECTION\s*B)[\s\-:]*SUPPLIES',
        'C': r'(?:SECTION\s*C)[\s\-:]*(?:DESCRIPTION|STATEMENT\s*OF\s*WORK|SOW|PWS)',
        'D': r'(?:SECTION\s*D)[\s\-:]*PACKAGING',
        'E': r'(?:SECTION\s*E)[\s\-:]*INSPECTION',
        'F': r'(?:SECTION\s*F)[\s\-:]*DELIVER',
        'G': r'(?:SECTION\s*G)[\s\-:]*CONTRACT\s*ADMIN',
        'H': r'(?:SECTION\s*H)[\s\-:]*SPECIAL',
        'I': r'(?:SECTION\s*I)[\s\-:]*CONTRACT\s*CLAUSES',
        'J': r'(?:SECTION\s*J)[\s\-:]*(?:LIST|ATTACH)',
        'K': r'(?:SECTION\s*K)[\s\-:]*REPRESENT',
        'L': r'(?:SECTION\s*L)[\s\-:]*INSTRUCT',
        'M': r'(?:SECTION\s*M)[\s\-:]*EVALUAT',
    }

    def __init__(self):
        self.pdf_parser = PDFParser()

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file with full traceability.

        Args:
            file_path: Path to the document

        Returns:
            ParsedDocument with text, sections, and source locations
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_path)
        elif ext == ".docx":
            return self._parse_docx(file_path)
        elif ext == ".xlsx":
            return self._parse_xlsx(file_path)
        else:
            raise ValueError(f"Unsupported document type: {ext}")

    def parse_bytes(
        self,
        content: bytes,
        filename: str,
        doc_type: Optional[DocumentType] = None
    ) -> ParsedDocument:
        """Parse document from bytes."""
        if doc_type is None:
            ext = Path(filename).suffix.lower()
            doc_type = {
                ".pdf": DocumentType.PDF,
                ".docx": DocumentType.DOCX,
                ".xlsx": DocumentType.XLSX,
            }.get(ext, DocumentType.UNKNOWN)

        if doc_type == DocumentType.PDF:
            return self._parse_pdf_bytes(content, filename)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx_bytes(content, filename)
        elif doc_type == DocumentType.XLSX:
            return self._parse_xlsx_bytes(content, filename)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")

    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF with full bbox traceability."""
        pdf_doc = self.pdf_parser.parse(file_path)

        # Extract texts with locations
        extracted_texts = []
        for page in pdf_doc.pages:
            for block in page.blocks:
                extracted_texts.append(ExtractedText(
                    text=block.text,
                    page=page.page_number,
                    bbox=block.bbox.to_dict(),
                    confidence=block.confidence,
                    context_before=block.context_before,
                    context_after=block.context_after,
                ))

        # Detect sections
        sections = self._detect_sections(pdf_doc.full_text, pdf_doc)

        return ParsedDocument(
            filename=pdf_doc.filename,
            doc_type=DocumentType.PDF,
            total_pages=pdf_doc.total_pages,
            full_text=pdf_doc.full_text,
            sections=sections,
            extracted_texts=extracted_texts,
            metadata=pdf_doc.metadata,
            _pdf_document=pdf_doc,
        )

    def _parse_pdf_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF from bytes."""
        pdf_doc = self.pdf_parser.parse_bytes(content, filename)

        extracted_texts = []
        for page in pdf_doc.pages:
            for block in page.blocks:
                extracted_texts.append(ExtractedText(
                    text=block.text,
                    page=page.page_number,
                    bbox=block.bbox.to_dict(),
                    confidence=block.confidence,
                    context_before=block.context_before,
                    context_after=block.context_after,
                ))

        sections = self._detect_sections(pdf_doc.full_text, pdf_doc)

        return ParsedDocument(
            filename=filename,
            doc_type=DocumentType.PDF,
            total_pages=pdf_doc.total_pages,
            full_text=pdf_doc.full_text,
            sections=sections,
            extracted_texts=extracted_texts,
            metadata=pdf_doc.metadata,
            _pdf_document=pdf_doc,
        )

    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX document."""
        from docx import Document

        doc = Document(file_path)
        extracted_texts = []
        full_text_parts = []
        current_page = 1  # DOCX doesn't have real pages, estimate

        for para in doc.paragraphs:
            if para.text.strip():
                extracted_texts.append(ExtractedText(
                    text=para.text,
                    page=current_page,
                    section=self._detect_section_from_text(para.text),
                ))
                full_text_parts.append(para.text)

                # Rough page estimation (50 lines per page)
                if len(extracted_texts) % 50 == 0:
                    current_page += 1

        # Process tables
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                extracted_texts.append(ExtractedText(
                    text=table_text,
                    page=current_page,
                ))
                full_text_parts.append(table_text)

        full_text = "\n\n".join(full_text_parts)
        sections = self._detect_sections_from_texts(extracted_texts)

        return ParsedDocument(
            filename=Path(file_path).name,
            doc_type=DocumentType.DOCX,
            total_pages=current_page,
            full_text=full_text,
            sections=sections,
            extracted_texts=extracted_texts,
            metadata={},
        )

    def _parse_docx_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX from bytes."""
        import io
        from docx import Document

        doc = Document(io.BytesIO(content))
        extracted_texts = []
        full_text_parts = []
        current_page = 1

        for para in doc.paragraphs:
            if para.text.strip():
                extracted_texts.append(ExtractedText(
                    text=para.text,
                    page=current_page,
                ))
                full_text_parts.append(para.text)
                if len(extracted_texts) % 50 == 0:
                    current_page += 1

        full_text = "\n\n".join(full_text_parts)

        return ParsedDocument(
            filename=filename,
            doc_type=DocumentType.DOCX,
            total_pages=current_page,
            full_text=full_text,
            sections=[],
            extracted_texts=extracted_texts,
            metadata={},
        )

    def _parse_xlsx(self, file_path: str) -> ParsedDocument:
        """Parse XLSX spreadsheet."""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        extracted_texts = []
        full_text_parts = []

        for sheet_idx, sheet in enumerate(wb.worksheets):
            sheet_text = f"=== Sheet: {sheet.title} ===\n"
            rows = []

            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value else "" for cell in row]
                if any(row_values):
                    rows.append("\t".join(row_values))

            sheet_content = "\n".join(rows)
            if sheet_content.strip():
                extracted_texts.append(ExtractedText(
                    text=sheet_content,
                    page=sheet_idx + 1,
                    section=sheet.title,
                ))
                full_text_parts.append(sheet_text + sheet_content)

        full_text = "\n\n".join(full_text_parts)

        return ParsedDocument(
            filename=Path(file_path).name,
            doc_type=DocumentType.XLSX,
            total_pages=len(wb.worksheets),
            full_text=full_text,
            sections=[],
            extracted_texts=extracted_texts,
            metadata={"sheets": [s.title for s in wb.worksheets]},
        )

    def _parse_xlsx_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse XLSX from bytes."""
        import io
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(content), data_only=True)
        extracted_texts = []
        full_text_parts = []

        for sheet_idx, sheet in enumerate(wb.worksheets):
            rows = []
            for row in sheet.iter_rows():
                row_values = [str(cell.value) if cell.value else "" for cell in row]
                if any(row_values):
                    rows.append("\t".join(row_values))

            sheet_content = "\n".join(rows)
            if sheet_content.strip():
                extracted_texts.append(ExtractedText(
                    text=sheet_content,
                    page=sheet_idx + 1,
                    section=sheet.title,
                ))
                full_text_parts.append(sheet_content)

        return ParsedDocument(
            filename=filename,
            doc_type=DocumentType.XLSX,
            total_pages=len(wb.worksheets),
            full_text="\n\n".join(full_text_parts),
            sections=[],
            extracted_texts=extracted_texts,
            metadata={},
        )

    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        return "\n".join(rows)

    def _detect_sections(
        self,
        full_text: str,
        pdf_doc: Optional[PDFDocument] = None
    ) -> List[DocumentSection]:
        """Detect FAR sections in the document."""
        sections = []

        for section_id, pattern in self.SECTION_PATTERNS.items():
            matches = list(re.finditer(pattern, full_text, re.IGNORECASE))
            for match in matches:
                # Try to find the page number
                page_num = 1
                if pdf_doc:
                    # Search for this text in the PDF to get page
                    results = pdf_doc.search_text(match.group(0))
                    if results:
                        page_num = results[0][0]

                sections.append(DocumentSection(
                    id=f"section_{section_id}",
                    title=f"Section {section_id}",
                    content="",  # Would need more parsing to extract full content
                    page_start=page_num,
                    page_end=page_num,
                    level=1,
                ))

        return sorted(sections, key=lambda s: s.page_start)

    def _detect_sections_from_texts(
        self,
        extracted_texts: List[ExtractedText]
    ) -> List[DocumentSection]:
        """Detect sections from extracted text blocks."""
        sections = []

        for ext in extracted_texts:
            section_id = self._detect_section_from_text(ext.text)
            if section_id:
                sections.append(DocumentSection(
                    id=section_id,
                    title=ext.text[:100],
                    content="",
                    page_start=ext.page,
                    page_end=ext.page,
                    level=1,
                ))

        return sections

    def _detect_section_from_text(self, text: str) -> Optional[str]:
        """Detect if text is a section header."""
        for section_id, pattern in self.SECTION_PATTERNS.items():
            if re.search(pattern, text, re.IGNORECASE):
                return f"section_{section_id}"
        return None


# Convenience function
def parse_document(file_path: str) -> ParsedDocument:
    """Parse a document with full traceability."""
    parser = DocumentParser()
    return parser.parse(file_path)
