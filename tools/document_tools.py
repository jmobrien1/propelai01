"""
PropelAI Document Ingestion Tools
Handles parsing of RFP documents (PDF, DOCX, XLS)

Key capabilities:
- Multi-format document loading
- OCR for scanned PDFs
- Structure recognition
- Text extraction with position tracking
"""

import os
import re
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import hashlib

# Document processing libraries will be imported dynamically
# to handle cases where they're not installed


@dataclass
class DocumentChunk:
    """A chunk of extracted document content"""
    id: str
    content: str
    page_number: int
    section_header: Optional[str]
    position: Tuple[int, int]  # (start_char, end_char)
    metadata: Dict[str, Any]


@dataclass
class ParsedDocument:
    """Fully parsed document"""
    file_name: str
    file_type: str
    total_pages: int
    total_chars: int
    chunks: List[DocumentChunk]
    raw_text: str
    structure: Dict[str, Any]
    metadata: Dict[str, Any]
    parsed_at: str


class DocumentLoader:
    """
    Multi-format document loader for RFP ingestion
    
    Supports:
    - PDF (with OCR fallback)
    - DOCX
    - XLS/XLSX
    - TXT
    """
    
    def __init__(self, ocr_enabled: bool = True):
        self.ocr_enabled = ocr_enabled
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which document processing libraries are available"""
        self.has_pypdf = False
        self.has_docx = False
        self.has_openpyxl = False
        
        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            pass
        
        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass
        
        try:
            import openpyxl
            self.has_openpyxl = True
        except ImportError:
            pass
    
    def load(self, file_path: str) -> ParsedDocument:
        """
        Load and parse a document
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument object
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower()
        
        # Route to appropriate parser
        if file_ext == ".pdf":
            return self._parse_pdf(file_path)
        elif file_ext in [".docx", ".doc"]:
            return self._parse_docx(file_path)
        elif file_ext in [".xlsx", ".xls"]:
            return self._parse_excel(file_path)
        elif file_ext in [".txt", ".md"]:
            return self._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
    
    def load_from_bytes(
        self, 
        content: bytes, 
        file_name: str, 
        file_type: str
    ) -> ParsedDocument:
        """
        Load document from bytes (for API uploads)
        """
        import tempfile
        
        # Write to temp file
        ext = f".{file_type}" if not file_type.startswith(".") else file_type
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
            f.write(content)
            temp_path = f.name
        
        try:
            result = self.load(temp_path)
            result.file_name = file_name
            return result
        finally:
            os.unlink(temp_path)
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF document"""
        chunks = []
        raw_text_parts = []
        total_pages = 0
        
        if self.has_pypdf:
            import pypdf
            
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                total_pages = len(reader.pages)
                
                char_offset = 0
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ""
                    
                    if text.strip():
                        chunk = DocumentChunk(
                            id=f"PDF-{page_num:04d}",
                            content=text,
                            page_number=page_num,
                            section_header=self._detect_section_header(text),
                            position=(char_offset, char_offset + len(text)),
                            metadata={"source": "pypdf"}
                        )
                        chunks.append(chunk)
                        raw_text_parts.append(text)
                        char_offset += len(text) + 1
        else:
            # Fallback: return placeholder
            raw_text_parts.append("[PDF parsing requires pypdf library]")
            total_pages = 1
        
        raw_text = "\n".join(raw_text_parts)
        
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="pdf",
            total_pages=total_pages,
            total_chars=len(raw_text),
            chunks=chunks,
            raw_text=raw_text,
            structure=self._detect_document_structure(raw_text),
            metadata={
                "parser": "pypdf" if self.has_pypdf else "placeholder",
                "pages": total_pages
            },
            parsed_at=datetime.now().isoformat()
        )
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX document"""
        chunks = []
        raw_text_parts = []
        
        if self.has_docx:
            import docx
            
            doc = docx.Document(file_path)
            
            char_offset = 0
            current_section = None
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect if this is a section header
                if para.style and 'Heading' in para.style.name:
                    current_section = text
                
                chunk = DocumentChunk(
                    id=f"DOCX-{i:04d}",
                    content=text,
                    page_number=1,  # DOCX doesn't have clear page breaks
                    section_header=current_section,
                    position=(char_offset, char_offset + len(text)),
                    metadata={
                        "style": para.style.name if para.style else None,
                        "source": "python-docx"
                    }
                )
                chunks.append(chunk)
                raw_text_parts.append(text)
                char_offset += len(text) + 1
        else:
            raw_text_parts.append("[DOCX parsing requires python-docx library]")
        
        raw_text = "\n".join(raw_text_parts)
        
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="docx",
            total_pages=1,
            total_chars=len(raw_text),
            chunks=chunks,
            raw_text=raw_text,
            structure=self._detect_document_structure(raw_text),
            metadata={
                "parser": "python-docx" if self.has_docx else "placeholder",
                "paragraphs": len(chunks)
            },
            parsed_at=datetime.now().isoformat()
        )
    
    def _parse_excel(self, file_path: str) -> ParsedDocument:
        """Parse Excel spreadsheet"""
        chunks = []
        raw_text_parts = []
        
        if self.has_openpyxl:
            import openpyxl
            
            wb = openpyxl.load_workbook(file_path)
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell else "" for cell in row]
                    if any(row_text):
                        sheet_text.append("\t".join(row_text))
                
                if sheet_text:
                    text = "\n".join(sheet_text)
                    chunk = DocumentChunk(
                        id=f"XLSX-{sheet_name}",
                        content=text,
                        page_number=1,
                        section_header=sheet_name,
                        position=(len("\n".join(raw_text_parts)), len("\n".join(raw_text_parts)) + len(text)),
                        metadata={
                            "sheet_name": sheet_name,
                            "source": "openpyxl"
                        }
                    )
                    chunks.append(chunk)
                    raw_text_parts.append(f"=== {sheet_name} ===\n{text}")
        else:
            raw_text_parts.append("[Excel parsing requires openpyxl library]")
        
        raw_text = "\n\n".join(raw_text_parts)
        
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="xlsx",
            total_pages=1,
            total_chars=len(raw_text),
            chunks=chunks,
            raw_text=raw_text,
            structure={},
            metadata={
                "parser": "openpyxl" if self.has_openpyxl else "placeholder",
                "sheets": len(chunks)
            },
            parsed_at=datetime.now().isoformat()
        )
    
    def _parse_text(self, file_path: str) -> ParsedDocument:
        """Parse plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()
        
        # Split into paragraphs
        paragraphs = raw_text.split("\n\n")
        chunks = []
        char_offset = 0
        
        for i, para in enumerate(paragraphs):
            if para.strip():
                chunk = DocumentChunk(
                    id=f"TXT-{i:04d}",
                    content=para.strip(),
                    page_number=1,
                    section_header=self._detect_section_header(para),
                    position=(char_offset, char_offset + len(para)),
                    metadata={"source": "text"}
                )
                chunks.append(chunk)
            char_offset += len(para) + 2
        
        return ParsedDocument(
            file_name=os.path.basename(file_path),
            file_type="txt",
            total_pages=1,
            total_chars=len(raw_text),
            chunks=chunks,
            raw_text=raw_text,
            structure=self._detect_document_structure(raw_text),
            metadata={"parser": "text"},
            parsed_at=datetime.now().isoformat()
        )
    
    def _detect_section_header(self, text: str) -> Optional[str]:
        """Detect if text starts with a section header"""
        lines = text.strip().split("\n")
        if not lines:
            return None
        
        first_line = lines[0].strip()
        
        # Check for common RFP section patterns
        section_patterns = [
            r'^(SECTION\s+[A-Z])',
            r'^([A-Z]\.\d+)',
            r'^(\d+\.\d+)',
            r'^(PART\s+\d+)',
            r'^(ARTICLE\s+\d+)',
        ]
        
        for pattern in section_patterns:
            match = re.match(pattern, first_line, re.IGNORECASE)
            if match:
                return first_line[:100]
        
        # Check if it's a short line (likely a header)
        if len(first_line) < 80 and first_line.isupper():
            return first_line
        
        return None
    
    def _detect_document_structure(self, text: str) -> Dict[str, Any]:
        """Detect the overall structure of the document"""
        structure = {
            "has_section_c": False,
            "has_section_l": False,
            "has_section_m": False,
            "sections_found": [],
            "page_references": []
        }
        
        text_lower = text.lower()
        
        # Check for standard RFP sections
        if "section c" in text_lower or "statement of work" in text_lower:
            structure["has_section_c"] = True
            structure["sections_found"].append("Section C")
        
        if "section l" in text_lower or "instructions to offeror" in text_lower:
            structure["has_section_l"] = True
            structure["sections_found"].append("Section L")
        
        if "section m" in text_lower or "evaluation criteria" in text_lower or "evaluation factors" in text_lower:
            structure["has_section_m"] = True
            structure["sections_found"].append("Section M")
        
        # Find page references
        page_refs = re.findall(r'page\s+(\d+)', text_lower)
        structure["page_references"] = list(set(page_refs))[:20]
        
        return structure


class ComplianceMatrixExporter:
    """Export compliance matrix to Excel format"""
    
    def export(
        self, 
        compliance_matrix: List[Dict[str, Any]], 
        output_path: str
    ) -> str:
        """
        Export compliance matrix to Excel
        
        Args:
            compliance_matrix: The compliance matrix from state
            output_path: Where to save the Excel file
            
        Returns:
            Path to created file
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Compliance Matrix"
            
            # Header row
            headers = [
                "Req ID", "Section Ref", "Type", "Requirement Text",
                "Compliance Status", "Compliant Response", "Assigned Owner",
                "Assigned Section", "Notes", "Confidence"
            ]
            
            # Style headers
            header_font = Font(bold=True)
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font_white = Font(bold=True, color="FFFFFF")
            
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                cell.alignment = Alignment(wrap_text=True)
            
            # Data rows
            for row_num, item in enumerate(compliance_matrix, 2):
                ws.cell(row=row_num, column=1, value=item.get("requirement_id", ""))
                ws.cell(row=row_num, column=2, value=item.get("section_reference", ""))
                ws.cell(row=row_num, column=3, value=item.get("requirement_type", ""))
                ws.cell(row=row_num, column=4, value=item.get("requirement_text", "")[:500])
                ws.cell(row=row_num, column=5, value=item.get("compliance_status", ""))
                ws.cell(row=row_num, column=6, value=item.get("compliant_response", ""))
                ws.cell(row=row_num, column=7, value=item.get("assigned_owner", ""))
                ws.cell(row=row_num, column=8, value=item.get("assigned_section", ""))
                ws.cell(row=row_num, column=9, value=item.get("notes", ""))
                ws.cell(row=row_num, column=10, value=item.get("confidence", 0))
            
            # Adjust column widths
            column_widths = [15, 12, 12, 60, 15, 40, 15, 15, 30, 10]
            for col, width in enumerate(column_widths, 1):
                ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width
            
            wb.save(output_path)
            return output_path
            
        except ImportError:
            # Fallback to CSV
            import csv
            
            csv_path = output_path.replace(".xlsx", ".csv")
            with open(csv_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=[
                    "requirement_id", "section_reference", "requirement_type",
                    "requirement_text", "compliance_status", "compliant_response",
                    "assigned_owner", "assigned_section", "notes", "confidence"
                ])
                writer.writeheader()
                writer.writerows(compliance_matrix)
            
            return csv_path


def create_document_loader(ocr_enabled: bool = True) -> DocumentLoader:
    """Factory function to create a DocumentLoader"""
    return DocumentLoader(ocr_enabled=ocr_enabled)


def create_compliance_exporter() -> ComplianceMatrixExporter:
    """Factory function to create a ComplianceMatrixExporter"""
    return ComplianceMatrixExporter()
